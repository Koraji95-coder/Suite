"""
Core rendering helpers for the Transmittal Builder.

These functions are extracted from the desktop (PyQt) implementation so they can
be reused in backend APIs without GUI dependencies.
"""

from __future__ import annotations

import os
import re
from typing import Dict, List, Optional

import pandas as pd
from docx import Document
from docx.table import Table

# ===================== FILENAME PARSING =====================
_H = r"[-–—]"
DOC_ID_RE = re.compile(rf"(R3P{_H}(\d+){_H}E(\d+){_H}(\d+))", re.IGNORECASE)


def extract_doc_meta(filename: str) -> dict:
    base = os.path.splitext(os.path.basename(filename))[0]
    m = DOC_ID_RE.search(base)
    if not m:
        return {"doc_no": "", "desc": base.strip(), "rev": ""}
    raw_doc = m.group(1)
    doc_no = re.sub(r"[–—]", "-", raw_doc).upper()

    # Prune R3P-PROJECT NUMBER- prefix, keep only E0-XXX, E1-XXX, etc.
    pruned_doc_no = re.sub(r"^R3P-\d+-", "", doc_no)

    remainder = base[m.end(1) :]
    remainder = re.sub(r"^[\s\-_–—:;|]+", "", remainder)
    desc_only = remainder.strip()
    return {"doc_no": pruned_doc_no, "desc": desc_only, "rev": ""}


def _norm_key_from_doc_no(doc_no: str) -> Optional[str]:
    m = re.search(r"E(\d+)[\-_–—](\d+)", doc_no or "", re.IGNORECASE)
    if not m:
        return None
    e = int(m.group(1))
    num = int(m.group(2))
    return f"E{e}-{num:04d}"


def build_drawing_index_items(paths: List[str]) -> List[dict]:
    items = [extract_doc_meta(f) for f in paths]
    sentinel = 10**12

    def parts(it: dict) -> tuple[int, int, int]:
        m = DOC_ID_RE.search(it.get("doc_no", "") or "")
        if not m:
            return (sentinel, sentinel, sentinel)
        proj = int(m.group(2))
        e = int(m.group(3))
        num = int(m.group(4))
        return (proj, e, num)

    items.sort(key=lambda it: (*parts(it), it["desc"].lower()))
    return items


# ===================== WORD HELPERS =====================

def _walk_paragraphs(doc: Document):
    """Yield every paragraph in the document, including nested tables and header/footer tables."""

    def iter_paragraphs_in_container(container):
        if hasattr(container, "paragraphs"):
            for p in container.paragraphs:
                yield p
        if hasattr(container, "tables"):
            for t in container.tables:
                for r in t.rows:
                    for c in r.cells:
                        yield from iter_paragraphs_in_container(c)

    # --- Body ---
    yield from iter_paragraphs_in_container(doc)

    # --- Headers & Footers ---
    for section in doc.sections:
        yield from iter_paragraphs_in_container(section.header)
        yield from iter_paragraphs_in_container(section.footer)


def _replace_text_in_run_sequence(runs, needle: str, replacement: str) -> bool:
    """Replace text that may span multiple runs."""
    if not needle:
        return False
    full = "".join(run.text for run in runs)
    idx = full.find(needle)
    if idx < 0:
        return False

    before, after = full[:idx], full[idx + len(needle) :]
    # Clear all runs
    for r in runs:
        r.text = ""

    # Distribute the new text across runs
    chunks = [before, replacement, after]
    ri = 0
    for chunk in chunks:
        if not chunk:
            continue
        if ri >= len(runs):
            runs[-1].text += chunk
        else:
            runs[ri].text = chunk
            ri += 1
    return True


def replace_text_everywhere(doc: Document, mapping: dict) -> dict:
    """Replace text placeholders throughout the document including headers, footers, and body."""
    replacements_made: Dict[str, str] = {}
    for p in _walk_paragraphs(doc):
        if not p.runs:
            continue
        full_text = "".join(r.text for r in p.runs)

        for key, val in mapping.items():
            if key in full_text:
                success = _replace_text_in_run_sequence(p.runs, key, val)
                if success:
                    replacements_made[key] = val

    return replacements_made


def set_checkbox_by_label(doc: Document, label: str, checked: bool) -> bool:
    """Robust checkbox toggler — handles nested tables, textboxes, and direct XML runs."""
    label_lc = (label or "").strip().lower()
    if not label_lc:
        return False

    def _find_glyphs_in_cell(cell):
        runs = []
        for p in cell.paragraphs:
            for r in p.runs:
                if any(sym in r.text for sym in ("☐", "☑", "☒")):
                    runs.append(r)
        for t in cell._element.xpath(".//w:t"):
            txt = t.text or ""
            if any(sym in txt for sym in ("☐", "☑", "☒")):
                runs.append(t)
        return runs

    def _scan_table(tbl: Table, level: int = 0) -> bool:
        for row in tbl.rows:
            label_col = None
            box_cols = []

            for c_idx, cell in enumerate(row.cells):
                text = " ".join(p.text for p in cell.paragraphs).strip().lower()
                if label_lc in text:
                    label_col = c_idx
                if any(sym in text for sym in ("☐", "☑", "☒")):
                    box_cols.append(c_idx)

            if label_col is not None:
                if any(b < label_col for b in box_cols):
                    chosen = max(b for b in box_cols if b < label_col)
                elif box_cols:
                    chosen = min(box_cols)
                else:
                    possible_cols = [label_col - 1, label_col]
                    chosen = None
                    for test_col in possible_cols:
                        if 0 <= test_col < len(row.cells):
                            deep_runs = _find_glyphs_in_cell(row.cells[test_col])
                            if deep_runs:
                                chosen = test_col
                                box_cols.append(chosen)
                                break

                if chosen is not None:
                    for r in _find_glyphs_in_cell(row.cells[chosen]):
                        if hasattr(r, "text"):
                            if checked:
                                r.text = r.text.replace("☐", "☒").replace("☑", "☒")
                            else:
                                r.text = r.text.replace("☑", "☐").replace("☒", "☐")
                            return True

            for cell in row.cells:
                for nested in cell.tables:
                    if _scan_table(nested, level + 1):
                        return True
        return False

    for tbl in doc.tables:
        if _scan_table(tbl):
            return True

    for section in doc.sections:
        if section.header:
            for tbl in section.header.tables:
                if _scan_table(tbl):
                    return True

    return False


def find_table_by_headers(doc: Document, headers: List[str]) -> Optional[Table]:
    for table in doc.tables:
        if not table.rows:
            continue
        first_row = table.rows[0]
        row_text = [cell.text.strip() for cell in first_row.cells]
        if all(h in row_text for h in headers):
            return table
    return None


def clear_table_body(table: Table) -> None:
    while len(table.rows) > 1:
        table._element.remove(table.rows[-1]._element)


def append_rows(table: Table, items: List[dict]) -> None:
    for item in items:
        row = table.add_row()
        row.cells[0].text = item.get("doc_no", "")
        row.cells[1].text = item.get("desc", "")
        row.cells[2].text = item.get("rev", "")


def fill_contacts_table(doc: Document, contacts: List[Dict[str, str]]) -> None:
    table = find_table_by_headers(doc, ["Name", "Company", "Email", "Phone"])
    if not table:
        return
    clear_table_body(table)
    for contact in contacts:
        if any(contact.values()):
            row = table.add_row()
            row.cells[0].text = contact.get("name", "")
            row.cells[1].text = contact.get("company", "")
            row.cells[2].text = contact.get("email", "")
            row.cells[3].text = contact.get("phone", "")


def fill_reference_table(doc: Document, reference_docs: List[dict]) -> None:
    """Fill the reference documents table."""
    reference_table = None
    possible_headers = [
        ["Reference Document No.", "Description", "Revision"],
        ["Reference No.", "Description", "Revision"],
        ["Reference", "Description", "Revision"],
        ["Document No.", "Description", "Revision"],
    ]

    for headers in possible_headers:
        reference_table = find_table_by_headers(doc, headers)
        if reference_table:
            break

    if not reference_table:
        return

    clear_table_body(reference_table)
    append_rows(reference_table, reference_docs)


def _norm_key_from_drawing_text(txt: str) -> Optional[str]:
    m = re.search(r"E\s*(\d+)\s*[-–—]?\s*(\d+)", txt)
    if not m:
        return None
    e = int(m.group(1))
    num = int(m.group(2))
    return f"E{e}-{num:04d}"


def load_revision_map_from_excel(xlsx_path: str) -> Dict[str, str]:
    xl = pd.ExcelFile(xlsx_path)
    for sheet in xl.sheet_names:
        df = xl.parse(sheet)
        if df.empty:
            continue
        cols = list(df.columns)
        dcol = rcol = None
        for c in cols:
            if "drawing" in str(c).lower() or "document" in str(c).lower():
                dcol = c
                break
        for c in cols:
            if "rev" in str(c).lower():
                rcol = c
                break
        if not (dcol and rcol):
            continue
        rev_map: Dict[str, str] = {}
        for _, row in df.iterrows():
            dv = row.get(dcol)
            if pd.isna(dv):
                continue
            key = _norm_key_from_drawing_text(str(dv).strip())
            if not key:
                continue
            rv = row.get(rcol)
            if pd.isna(rv):
                continue
            rev_map[key] = str(rv).strip()
        if rev_map:
            return rev_map
    raise RuntimeError(
        "Could not find drawing/revision columns in any sheet of the Excel file."
    )


def load_reference_documents_from_excel(xlsx_path: str) -> List[dict]:
    """Load reference documents from Excel file where column B contains 'Reference'."""
    xl = pd.ExcelFile(xlsx_path)
    reference_docs: List[dict] = []

    for sheet in xl.sheet_names:
        df = xl.parse(sheet)
        if df.empty:
            continue

        cols = list(df.columns)
        if len(cols) < 2:
            continue

        dcol = None
        for c in cols:
            if "drawing" in str(c).lower() or "document" in str(c).lower():
                dcol = c
                break

        if not dcol and len(cols) > 0:
            dcol = cols[0]

        bcol = cols[1] if len(cols) > 1 else None

        if not (dcol and bcol):
            continue

        rcol = None
        for c in cols:
            if "rev" in str(c).lower():
                rcol = c
                break

        for _, row in df.iterrows():
            bv = row.get(bcol)
            if pd.isna(bv):
                continue

            if "reference" in str(bv).lower():
                dv = row.get(dcol)
                if pd.isna(dv):
                    continue

                doc_text = str(dv).strip()
                if not doc_text:
                    continue

                rev = ""
                if rcol:
                    rv = row.get(rcol)
                    if not pd.isna(rv):
                        rev = str(rv).strip()

                doc_info = extract_doc_meta(doc_text)

                if not doc_info.get("doc_no"):
                    ref_match = re.match(
                        r"^([A-Z0-9\-]+)\s*[-–—]?\s*(.+)$", doc_text.strip()
                    )
                    if ref_match:
                        doc_info["doc_no"] = ref_match.group(1).strip()
                        doc_info["desc"] = ref_match.group(2).strip()
                    else:
                        doc_info["doc_no"] = ""
                        doc_info["desc"] = doc_text.strip()

                doc_info["rev"] = rev
                reference_docs.append(doc_info)

        if reference_docs:
            return reference_docs

    return reference_docs


# ===================== CORE RENDER =====================

def render_transmittal(
    template_path: str,
    documents_source: str,
    excel_path: str,
    fields: dict,
    checks: dict,
    contacts: List[Dict[str, str]],
    out_path: str,
    selected_files: Optional[List[str]] = None,
) -> str:
    if not os.path.isfile(template_path):
        raise FileNotFoundError("Template not found.")
    if not os.path.isfile(excel_path):
        raise FileNotFoundError("Drawing Index (Excel) not found.")

    if selected_files:
        documents_full = selected_files
        if not documents_full:
            raise ValueError("No document files selected.")
    else:
        if not os.path.isdir(documents_source):
            raise NotADirectoryError("Documents folder not found.")
        supported_extensions = (".pdf", ".cid")
        documents = [
            f for f in os.listdir(documents_source) if f.lower().endswith(supported_extensions)
        ]
        documents_full = [os.path.join(documents_source, f) for f in documents]
        if not documents_full:
            raise ValueError("No PDF or CID files found in the selected folder.")

    rev_map = load_revision_map_from_excel(excel_path)
    doc = Document(template_path)
    client_value = fields.get("client", "")
    mapping = {
        "<DATE>": fields.get("date", ""),
        "R3P-<PRJ#>": fields.get("job_num", ""),
        "XMTL-<###>": fields.get("transmittal_num", ""),
        "<CLIENT> - <SITE NAME>": client_value,
        "<PROJECT DESCRIPTION>": fields.get("project_desc", ""),
        "Andrew Simmons, P.E.": fields.get("from_name", ""),
        "Managing Partner": fields.get("from_title", ""),
        "e: andrew.simmons@root3power.com": f"e: {fields.get('from_email', '')}",
        "c: (713) 294-2003": f"c: {fields.get('from_phone', '')}",
        "TX FIRM #20290": fields.get("firm", ""),
    }

    replace_text_everywhere(doc, mapping)

    set_checkbox_by_label(doc, "PDF", checks["trans_pdf"])
    set_checkbox_by_label(doc, "CAD", checks["trans_cad"])
    set_checkbox_by_label(doc, "Originals", checks["trans_originals"])
    set_checkbox_by_label(doc, "Email", checks["via_email"])
    set_checkbox_by_label(doc, "FTP", checks["via_ftp"])
    set_checkbox_by_label(doc, "For Information Only", checks["ci_info"])
    set_checkbox_by_label(doc, "For Approval", checks["ci_approval"])
    set_checkbox_by_label(doc, "For Bid", checks["ci_bid"])
    set_checkbox_by_label(doc, "For Preliminary", checks["ci_preliminary"])
    set_checkbox_by_label(doc, "For Construction", checks["ci_const"])
    set_checkbox_by_label(doc, "For As-Built", checks["ci_asbuilt"])
    set_checkbox_by_label(doc, "For Fabrication", checks["ci_fab"])
    set_checkbox_by_label(doc, "For Record", checks["ci_record"])
    set_checkbox_by_label(doc, "For Reference", checks["ci_ref"])
    set_checkbox_by_label(doc, "Approved", checks["vr_approved"])
    set_checkbox_by_label(doc, "Approved as Noted", checks["vr_approved_noted"])
    set_checkbox_by_label(doc, "Rejected", checks["vr_rejected"])

    fill_contacts_table(doc, contacts)

    items = build_drawing_index_items(documents_full)
    for it in items:
        key = _norm_key_from_doc_no(it.get("doc_no", ""))
        it["rev"] = rev_map.get(key, "") if key else ""

    idx_table = find_table_by_headers(doc, ["Document No.", "Description", "Revision"])
    if not idx_table:
        raise RuntimeError(
            "Document Index table not found (expected header: Document No. | Description | Revision)."
        )
    clear_table_body(idx_table)
    append_rows(idx_table, items)

    try:
        reference_docs = load_reference_documents_from_excel(excel_path)
        if reference_docs:
            fill_reference_table(doc, reference_docs)
    except Exception:
        pass

    os.makedirs(os.path.dirname(out_path), exist_ok=True)
    doc.save(out_path)
    return out_path


def render_cid_transmittal(
    template_path: str,
    cid_folder: str,
    cid_index_data: List[Dict[str, str]],
    fields: dict,
    checks: dict,
    contacts: List[Dict[str, str]],
    out_path: str,
) -> str:
    if not os.path.isfile(template_path):
        raise FileNotFoundError("Template not found.")
    if not os.path.isdir(cid_folder):
        raise FileNotFoundError("CID folder not found.")
    if not cid_index_data:
        raise ValueError("No CID document data provided.")

    doc = Document(template_path)

    mapping = {
        "<DATE>": fields.get("date", ""),
        "R3P-<PRJ#>": fields.get("job_num", ""),
        "XMTL-<###>": fields.get("transmittal_num", ""),
        "<CLIENT> - <SITE NAME>": fields.get("client", ""),
        "<PROJECT DESCRIPTION>": fields.get("project_desc", ""),
        "Andrew Simmons, P.E.": fields.get("from_name", ""),
        "Managing Partner": fields.get("from_title", ""),
        "e: andrew.simmons@root3power.com": f"e: {fields.get('from_email', '')}",
        "c: (713) 294-2003": f"c: {fields.get('from_phone', '')}",
        "TX FIRM #20290": fields.get("firm", ""),
    }
    replace_text_everywhere(doc, mapping)

    set_checkbox_by_label(doc, "PDF", checks["trans_pdf"])
    set_checkbox_by_label(doc, "CAD", checks["trans_cad"])
    set_checkbox_by_label(doc, "Originals", checks["trans_originals"])
    set_checkbox_by_label(doc, "Email", checks["via_email"])
    set_checkbox_by_label(doc, "FTP", checks["via_ftp"])
    set_checkbox_by_label(doc, "For Information Only", checks["ci_info"])
    set_checkbox_by_label(doc, "For Approval", checks["ci_approval"])
    set_checkbox_by_label(doc, "For Bid", checks["ci_bid"])
    set_checkbox_by_label(doc, "For Preliminary", checks["ci_preliminary"])
    set_checkbox_by_label(doc, "For Construction", checks["ci_const"])
    set_checkbox_by_label(doc, "For As-Built", checks["ci_asbuilt"])
    set_checkbox_by_label(doc, "For Fabrication", checks["ci_fab"])
    set_checkbox_by_label(doc, "For Record", checks["ci_record"])
    set_checkbox_by_label(doc, "For Reference", checks["ci_ref"])
    set_checkbox_by_label(doc, "Approved", checks["vr_approved"])
    set_checkbox_by_label(doc, "Approved as Noted", checks["vr_approved_noted"])
    set_checkbox_by_label(doc, "Rejected", checks["vr_rejected"])

    fill_contacts_table(doc, contacts)
    populate_cid_document_index(doc, cid_index_data)

    os.makedirs(os.path.dirname(out_path), exist_ok=True)
    doc.save(out_path)
    return out_path


def populate_cid_document_index(doc: Document, cid_index_data: List[Dict[str, str]]) -> None:
    doc_table = None
    for table in doc.tables:
        if table.rows:
            header_row = table.rows[0]
            header_text = " ".join([cell.text.strip().lower() for cell in header_row.cells])
            if any(keyword in header_text for keyword in ["document", "drawing", "description", "revision"]):
                doc_table = table
                break

    if not doc_table:
        raise ValueError("Document index table not found in template.")

    while len(doc_table.rows) > 1:
        doc_table._element.remove(doc_table.rows[-1]._element)

    for i, cid_data in enumerate(cid_index_data):
        new_row = doc_table.add_row()
        if len(new_row.cells) >= 3:
            new_row.cells[0].text = str(i + 1)
            new_row.cells[1].text = cid_data["description"]
            new_row.cells[2].text = cid_data["revision"]
            if len(new_row.cells) >= 4:
                new_row.cells[3].text = cid_data["filename"]
