"""
Root3Power Transmittal Builder - Combined Version
Modern PyQt6 GUI with obsidian glass theme, SVG icons, and advanced styling
Includes all core logic from transmittal_builder_gui.py
"""

import sys
import os
import json
import platform
import subprocess
from datetime import datetime
from typing import Optional, Dict, List, Tuple

# PyQt6 imports
from PyQt6.QtWidgets import (
    QApplication, QMainWindow, QWidget, QVBoxLayout, QHBoxLayout,
    QGridLayout, QLabel, QPushButton, QLineEdit, QTextEdit, QCheckBox,
    QComboBox, QFrame, QScrollArea, QFileDialog, QMessageBox,
    QDialog, QPlainTextEdit,
    QGraphicsDropShadowEffect
)
from PyQt6.QtCore import Qt, QSize, QTimer
from PyQt6.QtGui import (
    QFont, QIcon, QColor, QPixmap
)
from PyQt6.QtSvgWidgets import QSvgWidget

# Configure high DPI scaling
import os
os.environ["QT_AUTO_SCREEN_SCALE_FACTOR"] = "1"
os.environ["QT_SCALE_FACTOR"] = "1"

# Logo will be loaded as PNG image

# Core transmittal builder functions (extracted from transmittal_builder_gui.py)
import re
import pandas as pd
from docx import Document
from docx.table import Table

# Import new modules
from core.config_system import ConfigManager
from widgets.inputs import DropLineEdit
from widgets.yaml_editor import YAMLEditorDialog
from utils.validation import (
    attach_line_validator,
    is_valid_email,
)

# Import email system
from emails.examples import send_bug_report_email, send_suggestion_email
from emails.templates import build_bug_email, build_suggestion_email

# ============================================================================
# CORE TRANSMITTAL BUILDER FUNCTIONS (from transmittal_builder_gui.py)
# ============================================================================

# ===================== FILENAME PARSING =====================
_H = r'[-–—]'
DOC_ID_RE = re.compile(rf'(R3P{_H}(\d+){_H}E(\d+){_H}(\d+))', re.IGNORECASE)

def extract_doc_meta(filename: str) -> dict:
    base = os.path.splitext(os.path.basename(filename))[0]
    m = DOC_ID_RE.search(base)
    if not m:
        return {'doc_no': '', 'desc': base.strip(), 'rev': ''}
    raw_doc = m.group(1)
    doc_no = re.sub(r'[–—]', '-', raw_doc).upper()

    # Prune R3P-PROJECT NUMBER- prefix, keep only E0-XXX, E1-XXX, etc.
    pruned_doc_no = re.sub(r'^R3P-\d+-', '', doc_no)

    remainder = base[m.end(1):]
    remainder = re.sub(r'^[\s\-_–—:;|]+', '', remainder)
    desc_only = remainder.strip()
    return {'doc_no': pruned_doc_no, 'desc': desc_only, 'rev': ''}

def _norm_key_from_doc_no(doc_no: str) -> Optional[str]:
    m = re.search(r'E(\d+)[\-_–—](\d+)', doc_no or '', re.IGNORECASE)
    if not m: return None
    e = int(m.group(1)); num = int(m.group(2))
    return f"E{e}-{num:04d}"

def build_drawing_index_items(paths: List[str]) -> List[dict]:
    items = [extract_doc_meta(f) for f in paths]
    SENTINEL = 10**12
    def parts(it):
        m = DOC_ID_RE.search(it.get('doc_no', '') or '')
        if not m: return (SENTINEL, SENTINEL, SENTINEL)
        proj = int(m.group(2)); e = int(m.group(3)); num = int(m.group(4))
        return (proj, e, num)
    items.sort(key=lambda it: (*parts(it), it['desc'].lower()))
    return items

# ===================== WORD HELPERS =====================
def _walk_paragraphs(doc: Document):
    """Yield every paragraph in the document, including nested tables and header/footer tables."""
    def iter_paragraphs_in_container(container):
        if hasattr(container, "paragraphs"):
            for p in container.paragraphs:
                yield p
        if hasattr(container, "tables"):
            for t in container.tables:
                for r in t.rows:
                    for c in r.cells:
                        yield from iter_paragraphs_in_container(c)

    # --- Body ---
    yield from iter_paragraphs_in_container(doc)

    # --- Headers & Footers ---
    for section in doc.sections:
        yield from iter_paragraphs_in_container(section.header)
        yield from iter_paragraphs_in_container(section.footer)


def _replace_text_in_run_sequence(runs, needle: str, replacement: str) -> bool:
    """Replace text that may span multiple runs."""
    if not needle:
        return False
    full = "".join(run.text for run in runs)
    idx = full.find(needle)
    if idx < 0:
        return False

    before, after = full[:idx], full[idx + len(needle):]
    # Clear all runs
    for r in runs:
        r.text = ""

    # Distribute the new text across runs
    chunks = [before, replacement, after]
    ri = 0
    for chunk in chunks:
        if not chunk:
            continue
        if ri >= len(runs):
            runs[-1].text += chunk
        else:
            runs[ri].text = chunk
            ri += 1
    return True

def replace_text_everywhere(doc, mapping: dict):
    """
    Replace text placeholders throughout the document including headers, footers, and body.
    """
    replacements_made = {}
    for p in _walk_paragraphs(doc):
        if not p.runs:
            continue
        full_text = "".join(r.text for r in p.runs)

        for key, val in mapping.items():
            if key in full_text:
                success = _replace_text_in_run_sequence(p.runs, key, val)
                if success:
                    replacements_made[key] = val

    return replacements_made


def set_checkbox_by_label(doc, label: str, checked: bool):
    """
    Robust checkbox toggler — handles nested tables, textboxes, and direct XML runs.
    """
    label_lc = (label or "").strip().lower()
    if not label_lc:
        return False

    def _flip(run):
        if checked:
            run.text = run.text.replace("☐", "☒").replace("☑", "☒")
        else:
            run.text = run.text.replace("☑", "☐").replace("☒", "☐")

    def _find_glyphs_in_cell(cell):
        """Return all runs and direct text nodes containing checkbox glyphs in this cell."""
        runs = []
        # normal runs
        for p in cell.paragraphs:
            for r in p.runs:
                if any(sym in r.text for sym in ("☐", "☑", "☒")):
                    runs.append(r)
        # deep XML text
        for t in cell._element.xpath(".//w:t"):
            txt = t.text or ""
            if any(sym in txt for sym in ("☐", "☑", "☒")):
                runs.append(t)
        return runs

    def _scan_table(tbl, t_idx, level=0):
        indent = "  " * level
        for r_idx, row in enumerate(tbl.rows):
            label_col = None
            box_cols = []

            for c_idx, cell in enumerate(row.cells):
                text = " ".join(p.text for p in cell.paragraphs).strip().lower()
                if label_lc in text:
                    label_col = c_idx
                if any(sym in text for sym in ("☐", "☑", "☒")):
                    box_cols.append(c_idx)

            if label_col is not None:
                # prefer a box to the left of label
                if any(b < label_col for b in box_cols):
                    chosen = max(b for b in box_cols if b < label_col)
                elif box_cols:
                    chosen = min(box_cols)
                else:
                    # fallback: even if no visible run, look deep in left neighbor
                    possible_cols = [label_col - 1, label_col]
                    for test_col in possible_cols:
                        if 0 <= test_col < len(row.cells):
                            deep_runs = _find_glyphs_in_cell(row.cells[test_col])
                            if deep_runs:
                                chosen = test_col
                                box_cols.append(chosen)
                                break
                    else:
                        chosen = None

                if chosen is not None:
                    for r in _find_glyphs_in_cell(row.cells[chosen]):
                        if hasattr(r, "text"):
                            if checked:
                                r.text = r.text.replace("☐", "☒").replace("☑", "☒")
                            else:
                                r.text = r.text.replace("☑", "☐").replace("☒", "☐")
                            return True

            # recurse into nested tables
            for cell in row.cells:
                for nested in cell.tables:
                    if _scan_table(nested, t_idx, level + 1):
                        return True
        return False

    # scan body
    for t_idx, tbl in enumerate(doc.tables):
        if _scan_table(tbl, t_idx):
            return True

    # scan headers
    for s_idx, section in enumerate(doc.sections):
        if section.header:
            for t_idx, tbl in enumerate(section.header.tables):
                if _scan_table(tbl, t_idx):
                    return True

    return False








def find_table_by_headers(doc: Document, headers: List[str]) -> Optional[Table]:
    for table in doc.tables:
        if not table.rows: continue
        first_row = table.rows[0]
        row_text = [cell.text.strip() for cell in first_row.cells]
        if all(h in row_text for h in headers):
            return table
    return None

def clear_table_body(table: Table):
    while len(table.rows) > 1:
        table._element.remove(table.rows[-1]._element)

def append_rows(table: Table, items: List[dict]):
    for item in items:
        row = table.add_row()
        row.cells[0].text = item.get('doc_no', '')
        row.cells[1].text = item.get('desc', '')
        row.cells[2].text = item.get('rev', '')

def fill_contacts_table(doc: Document, contacts: List[Dict[str, str]]):
    table = find_table_by_headers(doc, ["Name", "Company", "Email", "Phone"])
    if not table: return
    clear_table_body(table)
    for contact in contacts:
        if any(contact.values()):
            row = table.add_row()
            row.cells[0].text = contact.get('name', '')
            row.cells[1].text = contact.get('company', '')
            row.cells[2].text = contact.get('email', '')
            row.cells[3].text = contact.get('phone', '')

def fill_reference_table(doc: Document, reference_docs: List[dict]):
    """Fill the reference documents table."""
    # Look for reference table with various possible headers
    reference_table = None
    possible_headers = [
        ["Reference Document No.", "Description", "Revision"],
        ["Reference No.", "Description", "Revision"],
        ["Reference", "Description", "Revision"],
        ["Document No.", "Description", "Revision"]  # Fallback to main table format
    ]

    for headers in possible_headers:
        reference_table = find_table_by_headers(doc, headers)
        if reference_table:
            break

    if not reference_table:
        return

    clear_table_body(reference_table)
    append_rows(reference_table, reference_docs)


def _norm_key_from_drawing_text(txt: str) -> Optional[str]:
    m = re.search(r'E\s*(\d+)\s*[-–—]?\s*(\d+)', txt)
    if not m: return None
    e = int(m.group(1)); num = int(m.group(2))
    return f"E{e}-{num:04d}"

def load_revision_map_from_excel(xlsx_path: str) -> Dict[str, str]:
    xl = pd.ExcelFile(xlsx_path)
    for sheet in xl.sheet_names:
        df = xl.parse(sheet)
        if df.empty: continue
        cols = list(df.columns)
        dcol = rcol = None
        for c in cols:
            if 'drawing' in str(c).lower() or 'document' in str(c).lower():
                dcol = c; break
        for c in cols:
            if 'rev' in str(c).lower():
                rcol = c; break
        if not (dcol and rcol): continue
        rev_map = {}
        for _, row in df.iterrows():
            dv = row.get(dcol)
            if pd.isna(dv): continue
            key = _norm_key_from_drawing_text(str(dv).strip())
            if not key: continue
            rv = row.get(rcol)
            if pd.isna(rv): continue
            rev_map[key] = str(rv).strip()
        if rev_map:
            return rev_map
    raise RuntimeError("Could not find drawing/revision columns in any sheet of the Excel file.")

def load_reference_documents_from_excel(xlsx_path: str) -> List[dict]:
    """Load reference documents from Excel file where column B contains 'Reference'."""
    xl = pd.ExcelFile(xlsx_path)
    reference_docs = []

    for sheet in xl.sheet_names:
        df = xl.parse(sheet)
        if df.empty: continue

        # Get column names
        cols = list(df.columns)
        if len(cols) < 2: continue  # Need at least 2 columns

        # Find document column (usually column A)
        dcol = None
        for c in cols:
            if 'drawing' in str(c).lower() or 'document' in str(c).lower():
                dcol = c
                break

        # If no document column found, assume first column is documents
        if not dcol and len(cols) > 0:
            dcol = cols[0]

        # Column B is the second column
        bcol = cols[1] if len(cols) > 1 else None

        if not (dcol and bcol): continue

        # Find revision column
        rcol = None
        for c in cols:
            if 'rev' in str(c).lower():
                rcol = c
                break

        # Process rows where column B contains "Reference"
        for _, row in df.iterrows():
            bv = row.get(bcol)
            if pd.isna(bv): continue

            # Check if column B contains "Reference" (case insensitive)
            if 'reference' in str(bv).lower():
                dv = row.get(dcol)
                if pd.isna(dv): continue

                doc_text = str(dv).strip()
                if not doc_text: continue

                # Get revision if available
                rev = ''
                if rcol:
                    rv = row.get(rcol)
                    if not pd.isna(rv):
                        rev = str(rv).strip()

                # Extract document info - handle both R3P and non-R3P documents
                doc_info = extract_doc_meta(doc_text)

                # If extract_doc_meta didn't find a doc_no, try to extract it manually for reference docs
                if not doc_info.get('doc_no'):
                    # Look for patterns like "SPEC-001", "STD-001", etc.
                    import re
                    ref_match = re.match(r'^([A-Z0-9\-]+)\s*[-–—]?\s*(.+)$', doc_text.strip())
                    if ref_match:
                        doc_info['doc_no'] = ref_match.group(1).strip()
                        doc_info['desc'] = ref_match.group(2).strip()
                    else:
                        # If no pattern found, use the whole text as description
                        doc_info['doc_no'] = ''
                        doc_info['desc'] = doc_text.strip()

                doc_info['rev'] = rev
                reference_docs.append(doc_info)

        # If we found reference documents in this sheet, return them
        if reference_docs:
            return reference_docs

    # Return empty list if no reference documents found
    return reference_docs

# ===================== CORE RENDER =====================
def render_transmittal(template_path: str,
                       documents_source: str,  # Can be folder path or file list indicator
                       excel_path: str,
                       fields: dict,
                       checks: dict,
                       contacts: List[Dict[str, str]],
                       out_path: str,
                       selected_files: List[str] = None) -> str:
    if not os.path.isfile(template_path):
        raise FileNotFoundError("Template not found.")
    if not os.path.isfile(excel_path):
        raise FileNotFoundError("Drawing Index (Excel) not found.")

    # Determine document source and get file list
    if selected_files:
        # Individual files mode
        documents_full = selected_files
        if not documents_full:
            raise ValueError("No document files selected.")
    else:
        # Folder mode
        if not os.path.isdir(documents_source):
            raise NotADirectoryError("Documents folder not found.")
        supported_extensions = (".pdf", ".cid")
        documents = [f for f in os.listdir(documents_source) if f.lower().endswith(supported_extensions)]
        documents_full = [os.path.join(documents_source, f) for f in documents]
        if not documents_full:
            raise ValueError("No PDF or CID files found in the selected folder.")

    rev_map = load_revision_map_from_excel(excel_path)
    doc = Document(template_path)
    client_value = fields.get("client", "")
    mapping = {
        # Header placeholders (exactly as shown in template)
        "<DATE>": fields.get("date", ""),
        "R3P-<PRJ#>": fields.get("job_num", ""),
        "XMTL-<###>": fields.get("transmittal_num", ""),
        "<CLIENT> - <SITE NAME>": client_value,
        "<PROJECT DESCRIPTION>": fields.get("project_desc", ""),

        # FROM section placeholders (exactly as shown in template)
        "Andrew Simmons, P.E.": fields.get("from_name", ""),
        "Managing Partner": fields.get("from_title", ""),
        "e: andrew.simmons@root3power.com": f"e: {fields.get('from_email', '')}",
        "c: (713) 294-2003": f"c: {fields.get('from_phone', '')}",
        "TX FIRM #20290": fields.get("firm", ""),
    }

    replace_text_everywhere(doc, mapping)

    set_checkbox_by_label(doc, "PDF", checks["trans_pdf"])
    set_checkbox_by_label(doc, "CAD", checks["trans_cad"])
    set_checkbox_by_label(doc, "Originals", checks["trans_originals"])
    set_checkbox_by_label(doc, "Email", checks["via_email"])
    set_checkbox_by_label(doc, "FTP", checks["via_ftp"])
    set_checkbox_by_label(doc, "For Information Only", checks["ci_info"])
    set_checkbox_by_label(doc, "For Approval", checks["ci_approval"])
    set_checkbox_by_label(doc, "For Bid", checks["ci_bid"])
    set_checkbox_by_label(doc, "For Preliminary", checks["ci_preliminary"])
    set_checkbox_by_label(doc, "For Construction", checks["ci_const"])
    set_checkbox_by_label(doc, "For As-Built", checks["ci_asbuilt"])
    set_checkbox_by_label(doc, "For Fabrication", checks["ci_fab"])
    set_checkbox_by_label(doc, "For Record", checks["ci_record"])
    set_checkbox_by_label(doc, "For Reference", checks["ci_ref"])
    set_checkbox_by_label(doc, "Approved", checks["vr_approved"])
    set_checkbox_by_label(doc, "Approved as Noted", checks["vr_approved_noted"])
    set_checkbox_by_label(doc, "Rejected", checks["vr_rejected"])

    fill_contacts_table(doc, contacts)

    # Build document index from the file list
    items = build_drawing_index_items(documents_full)
    for it in items:
        key = _norm_key_from_doc_no(it.get('doc_no', ''))
        it['rev'] = rev_map.get(key, '') if key else ''

    idx_table = find_table_by_headers(doc, ["Document No.", "Description", "Revision"])
    if not idx_table:
        raise RuntimeError("Document Index table not found (expected header: Document No. | Description | Revision).")
    clear_table_body(idx_table)
    append_rows(idx_table, items)

    # Load and fill reference documents from Excel column B
    try:
        reference_docs = load_reference_documents_from_excel(excel_path)
        if reference_docs:
            fill_reference_table(doc, reference_docs)
    except Exception:
        pass  # Silently continue if reference documents can't be processed

    os.makedirs(os.path.dirname(out_path), exist_ok=True)
    doc.save(out_path)
    return out_path

def render_cid_transmittal(template_path: str,
                          cid_folder: str,
                          cid_index_data: List[Dict[str, str]],
                          fields: dict,
                          checks: dict,
                          contacts: List[Dict[str, str]],
                          out_path: str) -> str:
    """Render CID transmittal using CID files and generated index data."""
    if not os.path.isfile(template_path):
        raise FileNotFoundError("Template not found.")
    if not os.path.isdir(cid_folder):
        raise FileNotFoundError("CID folder not found.")
    if not cid_index_data:
        raise ValueError("No CID document data provided.")

    # Load template
    doc = Document(template_path)

    # Replace placeholders in document
    mapping = {
        # Header placeholders (exactly as shown in template)
        "<DATE>": fields.get("date", ""),
        "R3P-<PRJ#>": fields.get("job_num", ""),
        "XMTL-<###>": fields.get("transmittal_num", ""),
        "<CLIENT> - <SITE NAME>": fields.get("client", ""),
        "<PROJECT DESCRIPTION>": fields.get("project_desc", ""),

        # FROM section placeholders (exactly as shown in template)
        "Andrew Simmons, P.E.": fields.get("from_name", ""),
        "Managing Partner": fields.get("from_title", ""),
        "e: andrew.simmons@root3power.com": f"e: {fields.get('from_email', '')}",
        "c: (713) 294-2003": f"c: {fields.get('from_phone', '')}",
        "TX FIRM #20290": fields.get("firm", ""),
    }
    replace_text_everywhere(doc, mapping)

    # Toggle checkboxes
    set_checkbox_by_label(doc, "PDF", checks["trans_pdf"])
    set_checkbox_by_label(doc, "CAD", checks["trans_cad"])
    set_checkbox_by_label(doc, "Originals", checks["trans_originals"])
    set_checkbox_by_label(doc, "Email", checks["via_email"])
    set_checkbox_by_label(doc, "FTP", checks["via_ftp"])
    set_checkbox_by_label(doc, "For Information Only", checks["ci_info"])
    set_checkbox_by_label(doc, "For Approval", checks["ci_approval"])
    set_checkbox_by_label(doc, "For Bid", checks["ci_bid"])
    set_checkbox_by_label(doc, "For Preliminary", checks["ci_preliminary"])
    set_checkbox_by_label(doc, "For Construction", checks["ci_const"])
    set_checkbox_by_label(doc, "For As-Built", checks["ci_asbuilt"])
    set_checkbox_by_label(doc, "For Fabrication", checks["ci_fab"])
    set_checkbox_by_label(doc, "For Record", checks["ci_record"])
    set_checkbox_by_label(doc, "For Reference", checks["ci_ref"])
    set_checkbox_by_label(doc, "Approved", checks["vr_approved"])
    set_checkbox_by_label(doc, "Approved as Noted", checks["vr_approved_noted"])
    set_checkbox_by_label(doc, "Rejected", checks["vr_rejected"])

    # Populate TO table with contacts
    fill_contacts_table(doc, contacts)

    # Create document index from CID data
    populate_cid_document_index(doc, cid_index_data)

    # Save document
    os.makedirs(os.path.dirname(out_path), exist_ok=True)
    doc.save(out_path)
    return out_path

def populate_cid_document_index(doc, cid_index_data):
    """Populate document index table with CID data."""
    # Find the document index table
    doc_table = None
    for table in doc.tables:
        if table.rows and len(table.rows) > 0:
            # Check if this looks like a document index table
            header_row = table.rows[0]
            header_text = " ".join([cell.text.strip().lower() for cell in header_row.cells])

            # Look for common document index headers
            if any(keyword in header_text for keyword in ['document', 'drawing', 'description', 'revision']):
                doc_table = table
                break

    if not doc_table:
        raise ValueError("Document index table not found in template.")

    # Clear existing rows (except header)
    while len(doc_table.rows) > 1:
        doc_table._element.remove(doc_table.rows[-1]._element)

    # Add CID documents to table
    for i, cid_data in enumerate(cid_index_data):
        # Add new row
        new_row = doc_table.add_row()

        # Populate row cells (adjust based on your template structure)
        if len(new_row.cells) >= 3:
            # Assuming: Document Number | Description | Revision
            new_row.cells[0].text = str(i + 1)  # Sequential number
            new_row.cells[1].text = cid_data['description']
            new_row.cells[2].text = cid_data['revision']

            # Add filename as a comment or additional column if available
            if len(new_row.cells) >= 4:
                new_row.cells[3].text = cid_data['filename']

# ============================================================================
# CUSTOM WIDGETS
# ============================================================================

class NoScrollComboBox(QComboBox):
    """QComboBox that ignores mouse wheel events to prevent accidental changes."""

    def __init__(self, parent=None):
        super().__init__(parent)
        # Remove focus rectangle (dashed box)
        self.setStyleSheet("""
            QComboBox {
                outline: none;
            }
            QComboBox:focus {
                outline: none;
            }
        """)

    def wheelEvent(self, event):
        """Ignore wheel events."""
        event.ignore()

# Email functionality
import smtplib
from email.mime.multipart import MIMEMultipart
from email.mime.text import MIMEText

# ============================================================================
# CONSTANTS & CONFIGURATION
# ============================================================================

VERSION = "2.0"  # PyQt6 version

# Obsidian Glass Color Palette
OBSIDIAN_BG = "#0A0A0F"              # Deepest background
OBSIDIAN_CARD_DEEP = "#12141C"       # Deep card background
OBSIDIAN_CARD_MID = "#1A1D28"        # Mid-tone cards
OBSIDIAN_CARD_LIGHT = "#22253A"      # Lighter cards
OBSIDIAN_BORDER = "#2A2D42"          # Standard borders
OBSIDIAN_BORDER_GLOW = "#3A4A6A"     # Glowing borders

# Enhanced Accent Colors
OBSIDIAN_BLUE_DEEP = "#185FAC"       # Root3Power brand blue (R:24, G:95, B:172)
OBSIDIAN_BLUE_BRIGHT = "#4A9EFF"     # Bright blue
OBSIDIAN_BLUE_GLOW = "#6BB6FF"       # Glowing blue
ROOT3_BLUE = "#185FAC"               # Root3Power brand blue

# Text Colors
OBSIDIAN_TEXT_PRIMARY = "#E8EDF4"    # Bright white
OBSIDIAN_TEXT_ACCENT = "#185FAC"     # Root3Power brand blue
OBSIDIAN_TEXT_SECONDARY = "#8B92A0"  # Muted gray
OBSIDIAN_TEXT_DIM = "#5A5F6F"        # Very dim

# Status Colors
OBSIDIAN_SUCCESS = "#4ADE80"         # Bright green
OBSIDIAN_ERROR = "#FF5A6E"           # Bright red
OBSIDIAN_WARNING = "#FFB84D"         # Bright orange

# Initialize Config Manager (global instance)
config_manager = ConfigManager()
try:
    config_manager.ensure_ready()
    mail_creds = config_manager.mail_creds()
    SENDER_EMAIL = mail_creds.sender
    RECEIVER_EMAIL = mail_creds.default_receiver
    APP_PASSWORD = mail_creds.app_password
except Exception:
    # Fallback to empty values if config fails - user will need to configure
    SENDER_EMAIL = ""
    RECEIVER_EMAIL = ""
    APP_PASSWORD = ""

# ============================================================================
# UTILITY FUNCTIONS
# ============================================================================

def send_email_smtp(subject: str, body: str, cc_email: str = None) -> tuple[bool, str]:
    """
    Send email using SMTP (Gmail) - matches working implementation.

    Args:
        subject: Email subject
        body: Email body
        cc_email: Optional CC email address

    Returns:
        (success: bool, message: str)
    """
    import smtplib
    from email.mime.text import MIMEText
    from email.mime.multipart import MIMEMultipart
    from email.utils import formataddr

    try:
        # Get sender name from config
        sender_name = config_manager.get_sender_name() if config_manager else "Root3Power Transmittal Builder"

        # Create message
        msg = MIMEMultipart("alternative")
        msg['Subject'] = subject
        msg['From'] = formataddr((sender_name, SENDER_EMAIL))
        msg['To'] = RECEIVER_EMAIL

        # Add Reply-To if user provided email (for CC)
        if cc_email and cc_email.strip():
            msg['Reply-To'] = cc_email.strip()
            msg['Cc'] = cc_email.strip()

        # Escape HTML in body content
        def escape_html(text):
            return (text or "").replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")

        # Format body with better structure
        body_escaped = escape_html(body)

        # Attach body as HTML for better formatting
        html_body = f"""
<!DOCTYPE html>
<html>
<head>
    <meta charset="UTF-8">
    <meta name="viewport" content="width=device-width, initial-scale=1.0">
</head>
<body style="margin:0;padding:0;background:#0A0A0F;font-family:-apple-system,BlinkMacSystemFont,'Segoe UI',Roboto,Arial,sans-serif;">
    <div style="max-width:600px;margin:0 auto;background:#121212;padding:30px;border-radius:8px;">
        <!-- Header -->
        <div style="text-align:center;margin-bottom:25px;border-bottom:2px solid #185FAC;padding-bottom:20px;">
            <h1 style="color:#185FAC;margin:0 0 8px 0;font-size:28px;font-weight:600;">Root3Power LLC</h1>
            <p style="margin:0;font-size:16px;color:#8B92A0;font-weight:500;">Transmittal Builder</p>
        </div>

        <!-- Subject Badge -->
        <div style="background:#185FAC;color:#FFFFFF;padding:12px 20px;border-radius:6px;margin-bottom:20px;text-align:center;">
            <strong style="font-size:15px;">{escape_html(subject)}</strong>
        </div>

        <!-- Content -->
        <div style="background:#1E1E1E;padding:20px;border-radius:8px;border-left:4px solid #185FAC;">
            <div style="background:#2C2C2C;padding:15px;border-radius:6px;color:#E8EDF4;white-space:pre-wrap;font-size:14px;line-height:1.6;">
{body_escaped}
            </div>
        </div>

        <!-- Footer -->
        <div style="margin-top:25px;padding-top:15px;border-top:1px solid #2C2C2C;text-align:center;">
            <p style="margin:0;color:#8B92A0;font-size:12px;">
                <strong>App Version:</strong> {VERSION}
            </p>
            <p style="margin:8px 0 0 0;color:#5A5F6F;font-size:11px;">
                Root3Power Transmittal Builder | Electrical Engineering Solutions
            </p>
        </div>
    </div>
</body>
</html>
        """
        msg.attach(MIMEText(html_body, 'html'))

        # Prepare recipients list
        recipients = [RECEIVER_EMAIL]
        if cc_email and cc_email.strip():
            recipients.append(cc_email.strip())

        # Use SMTP_SSL directly (port 465) - same as working code
        with smtplib.SMTP_SSL('smtp.gmail.com', 465) as server:
            server.login(SENDER_EMAIL, APP_PASSWORD)
            server.sendmail(SENDER_EMAIL, recipients, msg.as_string())

        return True, "Email sent successfully!"

    except smtplib.SMTPAuthenticationError as e:
        error_msg = (
            "Gmail authentication failed.\n\n"
            "Please verify:\n"
            "1. 2-Step Verification is enabled on your Google account\n"
            "2. You're using an App Password (not your regular password)\n"
            "3. The App Password is correct (16 characters)\n\n"
            f"Error: {str(e)}"
        )
        return False, error_msg
    except smtplib.SMTPException as e:
        return False, f"SMTP error: {str(e)}"
    except Exception as e:
        return False, f"Failed to send email: {str(e)}"

def resource_path(relative_path):
    """Get absolute path to resource, works for dev and PyInstaller."""
    try:
        base_path = sys._MEIPASS
    except Exception:
        base_path = os.path.abspath(".")
    return os.path.join(base_path, relative_path)

LOGO_PATH = resource_path("assets/Root3Power_logo.png")
ICONS_PATH = resource_path("icons")

# ============================================================================
# OBSIDIAN GLASS STYLESHEET
# ============================================================================

OBSIDIAN_STYLESHEET = f"""
/* Main Window */
QMainWindow {{
    background-color: {OBSIDIAN_BG};
}}

/* Frames and Containers */
QFrame {{
    background-color: {OBSIDIAN_CARD_DEEP};
    border: 1px solid {OBSIDIAN_BORDER};
    border-radius: 8px;
}}

QFrame#sidebar {{
    background: qlineargradient(
        x1:0, y1:0, x2:0, y2:1,
        stop:0 {OBSIDIAN_CARD_DEEP},
        stop:1 {OBSIDIAN_CARD_MID}
    );
    border-right: 2px solid {OBSIDIAN_BORDER_GLOW};
    border-radius: 0px;
}}

QFrame#mainContent {{
    background-color: {OBSIDIAN_BG};
    border: 2px solid {OBSIDIAN_BLUE_DEEP};
    border-radius: 12px;
}}

QFrame#glassCard {{
    background: qlineargradient(
        x1:0, y1:0, x2:0, y2:1,
        stop:0 {OBSIDIAN_CARD_MID},
        stop:1 {OBSIDIAN_CARD_DEEP}
    );
    border: 1px solid {OBSIDIAN_BORDER};
    border-radius: 10px;
    padding: 15px;
}}

/* Labels */
QLabel {{
    color: {OBSIDIAN_TEXT_PRIMARY};
    background: transparent;
    border: none;
}}

QLabel#sectionHeader {{
    color: {OBSIDIAN_TEXT_ACCENT};
    font-size: 14px;
    font-weight: bold;
}}

QLabel#title {{
    color: {OBSIDIAN_TEXT_ACCENT};
    font-size: 18px;
    font-weight: bold;
}}

QLabel#subtitle {{
    color: {OBSIDIAN_TEXT_SECONDARY};
    font-size: 11px;
}}

/* Buttons */
QPushButton {{
    background: qlineargradient(
        x1:0, y1:0, x2:0, y2:1,
        stop:0 {OBSIDIAN_BLUE_DEEP},
        stop:0.5 {OBSIDIAN_CARD_MID},
        stop:1 {OBSIDIAN_CARD_DEEP}
    );
    color: {OBSIDIAN_TEXT_PRIMARY};
    border: 2px solid {OBSIDIAN_BORDER};
    border-radius: 8px;
    padding: 8px 16px;
    font-size: 13px;
    font-weight: bold;
}}

QPushButton:hover {{
    background: qlineargradient(
        x1:0, y1:0, x2:0, y2:1,
        stop:0 {OBSIDIAN_BLUE_BRIGHT},
        stop:0.5 {OBSIDIAN_BLUE_DEEP},
        stop:1 {OBSIDIAN_CARD_MID}
    );
    border: 2px solid {OBSIDIAN_BLUE_GLOW};
}}

QPushButton:pressed {{
    background: qlineargradient(
        x1:0, y1:0, x2:0, y2:1,
        stop:0 {OBSIDIAN_CARD_DEEP},
        stop:0.5 {OBSIDIAN_CARD_MID},
        stop:1 {OBSIDIAN_BLUE_DEEP}
    );
    border: 2px solid {OBSIDIAN_BORDER};
}}

/* Preview and Generate buttons now use default button styling for consistency */

/* Support buttons now use default button styling for consistency */

/* Input Fields */
QLineEdit, QTextEdit, QPlainTextEdit {{
    background: qlineargradient(
        x1:0, y1:0, x2:0, y2:1,
        stop:0 {OBSIDIAN_CARD_DEEP},
        stop:1 {OBSIDIAN_CARD_MID}
    );
    color: {OBSIDIAN_TEXT_PRIMARY};
    border: 2px solid {OBSIDIAN_BORDER};
    border-radius: 6px;
    padding: 6px;
    font-size: 12px;
}}

QLineEdit:hover, QTextEdit:hover, QPlainTextEdit:hover {{
    border: 2px solid {OBSIDIAN_BORDER_GLOW};
}}

QLineEdit:focus, QTextEdit:focus, QPlainTextEdit:focus {{
    border: 2px solid {OBSIDIAN_BLUE_GLOW};
    background: qlineargradient(
        x1:0, y1:0, x2:0, y2:1,
        stop:0 {OBSIDIAN_CARD_MID},
        stop:1 {OBSIDIAN_CARD_LIGHT}
    );
}}

/* ComboBox */
QComboBox {{
    background: qlineargradient(
        x1:0, y1:0, x2:0, y2:1,
        stop:0 {OBSIDIAN_CARD_DEEP},
        stop:1 {OBSIDIAN_CARD_MID}
    );
    color: {OBSIDIAN_TEXT_PRIMARY};
    border: 2px solid {OBSIDIAN_BORDER};
    border-radius: 6px;
    padding: 6px;
    padding-right: 25px;
    font-size: 12px;
}}

QComboBox:hover {{
    border: 2px solid {OBSIDIAN_BORDER_GLOW};
    background: qlineargradient(
        x1:0, y1:0, x2:0, y2:1,
        stop:0 {OBSIDIAN_CARD_MID},
        stop:1 {OBSIDIAN_CARD_LIGHT}
    );
}}

QComboBox:focus {{
    border: 2px solid {OBSIDIAN_BLUE_GLOW};
}}

QComboBox::drop-down {{
    border: none;
    width: 25px;
    background: transparent;
}}

QComboBox::down-arrow {{
    image: none;
    border-left: 5px solid transparent;
    border-right: 5px solid transparent;
    border-top: 6px solid {OBSIDIAN_TEXT_ACCENT};
    margin-right: 5px;
}}

QComboBox QAbstractItemView {{
    background: qlineargradient(
        x1:0, y1:0, x2:0, y2:1,
        stop:0 {OBSIDIAN_CARD_MID},
        stop:1 {OBSIDIAN_CARD_DEEP}
    );
    color: {OBSIDIAN_TEXT_PRIMARY};
    selection-background-color: {OBSIDIAN_BLUE_DEEP};
    selection-color: {OBSIDIAN_TEXT_PRIMARY};
    border: 2px solid {OBSIDIAN_BORDER_GLOW};
    border-radius: 6px;
    padding: 4px;
}}

QComboBox QAbstractItemView::item {{
    padding: 6px;
    border-radius: 4px;
}}

QComboBox QAbstractItemView::item:hover {{
    background-color: {OBSIDIAN_BLUE_DEEP};
}}

/* CheckBox */
QCheckBox {{
    color: {OBSIDIAN_TEXT_PRIMARY};
    spacing: 8px;
}}

QCheckBox::indicator {{
    width: 20px;
    height: 20px;
    border: 2px solid {OBSIDIAN_BORDER};
    border-radius: 5px;
    background: qlineargradient(
        x1:0, y1:0, x2:0, y2:1,
        stop:0 {OBSIDIAN_CARD_DEEP},
        stop:1 {OBSIDIAN_CARD_MID}
    );
}}

QCheckBox::indicator:hover {{
    border: 2px solid {OBSIDIAN_BORDER_GLOW};
    background: qlineargradient(
        x1:0, y1:0, x2:0, y2:1,
        stop:0 {OBSIDIAN_CARD_MID},
        stop:1 {OBSIDIAN_CARD_LIGHT}
    );
}}

QCheckBox::indicator:checked {{
    background: qlineargradient(
        x1:0, y1:0, x2:0, y2:1,
        stop:0 {OBSIDIAN_BLUE_BRIGHT},
        stop:1 {OBSIDIAN_BLUE_DEEP}
    );
    border: 2px solid {OBSIDIAN_BLUE_GLOW};
}}

QCheckBox::indicator:checked:hover {{
    background: qlineargradient(
        x1:0, y1:0, x2:0, y2:1,
        stop:0 {OBSIDIAN_BLUE_GLOW},
        stop:1 {OBSIDIAN_BLUE_BRIGHT}
    );
}}

/* ScrollArea */
QScrollArea {{
    background: transparent;
    border: none;
}}

QScrollBar:vertical {{
    background: qlineargradient(
        x1:0, y1:0, x2:1, y2:0,
        stop:0 {OBSIDIAN_CARD_DEEP},
        stop:1 {OBSIDIAN_CARD_MID}
    );
    width: 14px;
    border-radius: 7px;
    margin: 2px;
}}

QScrollBar::handle:vertical {{
    background: qlineargradient(
        x1:0, y1:0, x2:1, y2:0,
        stop:0 {OBSIDIAN_BORDER_GLOW},
        stop:1 {OBSIDIAN_BLUE_DEEP}
    );
    border-radius: 6px;
    min-height: 30px;
    margin: 2px;
}}

QScrollBar::handle:vertical:hover {{
    background: qlineargradient(
        x1:0, y1:0, x2:1, y2:0,
        stop:0 {OBSIDIAN_BLUE_GLOW},
        stop:1 {OBSIDIAN_BLUE_BRIGHT}
    );
}}

QScrollBar::add-line:vertical, QScrollBar::sub-line:vertical {{
    height: 0px;
}}

QScrollBar:horizontal {{
    background: qlineargradient(
        x1:0, y1:0, x2:0, y2:1,
        stop:0 {OBSIDIAN_CARD_DEEP},
        stop:1 {OBSIDIAN_CARD_MID}
    );
    height: 14px;
    border-radius: 7px;
    margin: 2px;
}}

QScrollBar::handle:horizontal {{
    background: qlineargradient(
        x1:0, y1:0, x2:0, y2:1,
        stop:0 {OBSIDIAN_BORDER_GLOW},
        stop:1 {OBSIDIAN_BLUE_DEEP}
    );
    border-radius: 6px;
    min-width: 30px;
    margin: 2px;
}}

QScrollBar::handle:horizontal:hover {{
    background: qlineargradient(
        x1:0, y1:0, x2:0, y2:1,
        stop:0 {OBSIDIAN_BLUE_GLOW},
        stop:1 {OBSIDIAN_BLUE_BRIGHT}
    );
}}

QScrollBar::add-line:horizontal, QScrollBar::sub-line:horizontal {{
    width: 0px;
}}

/* Terminal/Log TextEdit */
QTextEdit#terminal {{
    background-color: {OBSIDIAN_CARD_MID};
    color: {OBSIDIAN_TEXT_PRIMARY};
    border: 1px solid {OBSIDIAN_BORDER};
    border-radius: 8px;
    font-family: 'Consolas', 'Courier New', monospace;
    font-size: 10px;
}}
"""

# ============================================================================
# ICON MANAGER
# ============================================================================

class IconManager:
    """Manages loading and caching of Lucide SVG icons."""

    def __init__(self):
        self.icons_cache = {}
        self.icons_path = resource_path("icons")

    def get_icon(self, icon_name: str, size: int = 24, color: str = None) -> QIcon:
        """
        Get a QIcon from an SVG file.

        Args:
            icon_name: Name of the icon file (e.g., 'bug.svg' or 'bug')
            size: Size of the icon in pixels
            color: Optional color to tint the icon (hex color string)

        Returns:
            QIcon object
        """
        # Add .svg extension if not present
        if not icon_name.endswith('.svg'):
            icon_name = f"{icon_name}.svg"

        # Create cache key
        cache_key = f"{icon_name}_{size}_{color}"

        # Return cached icon if available
        if cache_key in self.icons_cache:
            return self.icons_cache[cache_key]

        # Load SVG file
        icon_path = os.path.join(self.icons_path, icon_name)

        if not os.path.exists(icon_path):
            return QIcon()

        # Read SVG content
        with open(icon_path, 'r', encoding='utf-8') as f:
            svg_content = f.read()

        # Apply color if specified
        if color:
            # Replace stroke color in SVG
            svg_content = svg_content.replace('stroke="currentColor"', f'stroke="{color}"')
            svg_content = svg_content.replace('fill="none"', f'fill="none"')

        # Create QIcon from SVG
        icon = QIcon(icon_path)

        # Cache the icon
        self.icons_cache[cache_key] = icon

        return icon

    def get_pixmap(self, icon_name: str, size: int = 24, color: str = None) -> QPixmap:
        """
        Get a QPixmap from an SVG file.

        Args:
            icon_name: Name of the icon file
            size: Size of the pixmap in pixels
            color: Optional color to tint the icon

        Returns:
            QPixmap object
        """
        icon = self.get_icon(icon_name, size, color)
        return icon.pixmap(size, size)

    def create_svg_widget(self, icon_name: str, size: int = 24) -> QSvgWidget:
        """
        Create a QSvgWidget for displaying an SVG icon.

        Args:
            icon_name: Name of the icon file
            size: Size of the widget in pixels

        Returns:
            QSvgWidget object
        """
        # Add .svg extension if not present
        if not icon_name.endswith('.svg'):
            icon_name = f"{icon_name}.svg"

        icon_path = os.path.join(self.icons_path, icon_name)

        if not os.path.exists(icon_path):
            return QSvgWidget()

        svg_widget = QSvgWidget(icon_path)
        svg_widget.setFixedSize(size, size)

        return svg_widget

# ============================================================================
# UTILITY FUNCTIONS
# ============================================================================

def send_email(subject: str, body: str, is_html: bool = False) -> bool:
    """Send email with the given subject and body."""
    try:
        msg = MIMEMultipart('alternative')
        msg['Subject'] = subject
        msg['From'] = SENDER_EMAIL
        msg['To'] = RECEIVER_EMAIL
        
        if is_html:
            plain_text = body.replace('<br>', '\n').replace('<b>', '').replace('</b>', '')
            msg.attach(MIMEText(plain_text, 'plain'))
            msg.attach(MIMEText(body, 'html'))
        else:
            msg.attach(MIMEText(body, 'plain'))
        
        with smtplib.SMTP_SSL('smtp.gmail.com', 465) as server:
            server.login(SENDER_EMAIL, APP_PASSWORD)
            server.sendmail(SENDER_EMAIL, RECEIVER_EMAIL, msg.as_string())
        
        return True
    except Exception:
        return False

# ============================================================================
# MAIN APPLICATION CLASS
# ============================================================================

class TransmittalBuilderPyQt6(QMainWindow):
    """Main application window for Transmittal Builder using PyQt6."""
    
    def __init__(self):
        super().__init__()

        # Initialize icon manager
        self.icon_manager = IconManager()

        # Initialize data structures
        self.contacts = []
        self.project_data = {}
        self.auto_save_timer = None
        self.last_template_dir = ""
        self.last_index_dir = ""
        self.last_output_dir = ""

        # Pristine validation mode (no red borders until user takes action)
        self._show_validation = False
        self._validator_runs = []  # Store validator repaint functions

        # Setup UI (this creates the log widget and contacts)
        self.init_ui()

        # Log startup message only (no contact log)
        self.log_message("Transmittal Builder started", "info")

        self.setup_auto_save()

        # Setup live validation
        self.setup_validation()

        # Load previous session
        self.load_autosave()

    def resizeEvent(self, event):
        """Handle window resize events for responsive design."""
        super().resizeEvent(event)

        # Adjust layout based on window size
        if hasattr(self, 'sidebar_frame') and hasattr(self, 'main_content_frame'):
            window_width = self.width()

            # If window becomes too narrow, adjust sidebar
            if window_width < 1000:
                # Make sidebar narrower on small windows
                if hasattr(self, 'sidebar_frame'):
                    self.sidebar_frame.setFixedWidth(320)
            elif window_width < 1200:
                # Medium width
                if hasattr(self, 'sidebar_frame'):
                    self.sidebar_frame.setFixedWidth(380)
            else:
                # Full width - restore normal sidebar
                if hasattr(self, 'sidebar_frame'):
                    if hasattr(self, 'screen_width'):
                        if self.screen_width <= 1366:
                            self.sidebar_frame.setFixedWidth(380)
                        elif self.screen_width <= 1920:
                            self.sidebar_frame.setFixedWidth(420)
                        else:
                            self.sidebar_frame.setFixedWidth(450)
                    else:
                        self.sidebar_frame.setFixedWidth(450)

    def on_screen_changed(self):
        """Handle screen changes (moving between monitors, resolution changes)."""
        # Get current screen the window is on
        current_screen = QApplication.screenAt(self.pos())
        if not current_screen:
            current_screen = QApplication.primaryScreen()

        new_geometry = current_screen.availableGeometry()
        new_width = new_geometry.width()
        new_height = new_geometry.height()

        # Check if we've moved to a significantly different screen size
        if hasattr(self, 'screen_width'):
            width_change = abs(new_width - self.screen_width) / self.screen_width
            height_change = abs(new_height - self.screen_height) / self.screen_height

            # If screen size changed by more than 20%, adjust the window
            if width_change > 0.2 or height_change > 0.2:
                self.log_message(f"Screen change detected: {new_width}x{new_height}", "info")
                self.adjust_for_screen_size(new_width, new_height, new_geometry)

    def adjust_for_screen_size(self, screen_width, screen_height, screen_geometry):
        """Adjust window size and layout for new screen dimensions."""
        # Update stored screen info
        self.screen_width = screen_width
        self.screen_height = screen_height

        # Calculate new window size based on screen category (slightly larger for better button fit)
        if screen_width <= 1366:  # Small screens
            width = min(1150, int(screen_width * 0.88))
            height = min(750, int(screen_height * 0.82))
            sidebar_width = 380
        elif screen_width <= 1920:  # Medium screens
            width = min(1250, int(screen_width * 0.78))
            height = min(850, int(screen_height * 0.78))
            sidebar_width = 420
        else:  # Large screens
            width = min(1450, int(screen_width * 0.68))
            height = min(950, int(screen_height * 0.72))
            sidebar_width = 450

        # Center the window on new screen
        x = screen_geometry.x() + (screen_width - width) // 2
        y = screen_geometry.y() + (screen_height - height) // 2

        # Apply new geometry
        self.setGeometry(x, y, width, height)

        # Update sidebar width
        if hasattr(self, 'sidebar_frame'):
            self.sidebar_frame.setFixedWidth(sidebar_width)

    def moveEvent(self, event):
        """Handle window move events to detect screen changes."""
        super().moveEvent(event)

        # Check if we've moved to a different screen
        if hasattr(self, 'current_screen'):
            new_screen = QApplication.screenAt(self.pos())
            if new_screen and new_screen != self.current_screen:
                self.current_screen = new_screen
                # Trigger screen change detection with a small delay
                QTimer.singleShot(100, self.on_screen_changed)

    def init_ui(self):
        """Initialize the user interface."""
        # Window configuration
        self.setWindowTitle("Root3Power Transmittal Builder")

        # Get screen geometry for responsive sizing
        screen = QApplication.primaryScreen()
        screen_geometry = screen.availableGeometry()
        screen_width = screen_geometry.width()
        screen_height = screen_geometry.height()

        # Responsive sizing based on screen size (slightly larger for better button fit)
        if screen_width <= 1366:  # Small screens (laptops, small monitors)
            width = min(1150, int(screen_width * 0.88))
            height = min(750, int(screen_height * 0.82))
            self.setMinimumSize(850, 550)
        elif screen_width <= 1920:  # Medium screens (standard monitors)
            width = min(1250, int(screen_width * 0.78))
            height = min(850, int(screen_height * 0.78))
            self.setMinimumSize(950, 600)
        else:  # Large screens (4K, ultrawide)
            width = min(1450, int(screen_width * 0.68))
            height = min(950, int(screen_height * 0.72))
            self.setMinimumSize(1050, 650)

        # Center the window
        x = (screen_width - width) // 2
        y = (screen_height - height) // 2

        self.setGeometry(x, y, width, height)

        # Store screen info for responsive adjustments
        self.screen_width = screen_width
        self.screen_height = screen_height
        self.current_screen = screen

        # Connect to screen change detection
        app = QApplication.instance()
        app.screenAdded.connect(self.on_screen_changed)
        app.screenRemoved.connect(self.on_screen_changed)
        app.primaryScreenChanged.connect(self.on_screen_changed)

        # Set window icon (PNG logo)
        if os.path.exists(LOGO_PATH):
            self.setWindowIcon(QIcon(LOGO_PATH))

        # Apply obsidian stylesheet
        self.setStyleSheet(OBSIDIAN_STYLESHEET)
        
        # Create central widget and main layout
        central_widget = QWidget()
        self.setCentralWidget(central_widget)
        
        main_layout = QHBoxLayout(central_widget)
        main_layout.setContentsMargins(0, 0, 0, 0)
        main_layout.setSpacing(0)
        
        # Create sidebar and main content
        self.create_sidebar()
        self.create_main_content()
        
        # Add to main layout
        main_layout.addWidget(self.sidebar_frame)
        main_layout.addWidget(self.main_content_frame, 1)  # Stretch factor 1
    
    def create_sidebar(self):
        """Create the left sidebar with controls and activity log."""
        self.sidebar_frame = QFrame()
        self.sidebar_frame.setObjectName("sidebar")

        # Responsive sidebar width based on screen size
        if hasattr(self, 'screen_width'):
            if self.screen_width <= 1366:  # Small screens
                sidebar_width = 380
            elif self.screen_width <= 1920:  # Medium screens
                sidebar_width = 420
            else:  # Large screens
                sidebar_width = 450
        else:
            sidebar_width = 450  # Default fallback

        self.sidebar_frame.setFixedWidth(sidebar_width)

        # Add subtle glow to sidebar
        self.add_drop_shadow(self.sidebar_frame, blur_radius=30, offset_x=5, offset_y=0,
                            color=QColor(OBSIDIAN_BLUE_DEEP))
        
        sidebar_layout = QVBoxLayout(self.sidebar_frame)
        sidebar_layout.setContentsMargins(20, 20, 20, 20)
        sidebar_layout.setSpacing(15)
        
        # Logo and title (placeholder for now)
        self.create_header(sidebar_layout)
        
        # Action buttons
        self.create_action_buttons(sidebar_layout)
        
        # Status indicator
        self.create_status_indicator(sidebar_layout)
        
        # Activity log
        self.create_activity_log(sidebar_layout)
        
        # Support buttons
        self.create_support_buttons(sidebar_layout)
        
        # Version footer
        self.create_version_footer(sidebar_layout)
    
    def create_header(self, layout):
        """Create logo and title header with animated 3D logo."""
        header_container = QWidget()
        header_layout = QVBoxLayout(header_container)
        header_layout.setContentsMargins(10, 15, 10, 10)
        header_layout.setSpacing(5)

        # PNG Logo
        logo_label = QLabel()
        if os.path.exists(LOGO_PATH):
            pixmap = QPixmap(LOGO_PATH)
            # Scale the logo to a reasonable size
            scaled_pixmap = pixmap.scaled(80, 80, Qt.AspectRatioMode.KeepAspectRatio, Qt.TransformationMode.SmoothTransformation)
            logo_label.setPixmap(scaled_pixmap)
        else:
            # Fallback text if logo not found
            logo_label.setText("R3P")
            logo_label.setStyleSheet(f"""
                font-size: 24px;
                font-weight: bold;
                color: {OBSIDIAN_TEXT_ACCENT};
                border: 2px solid {OBSIDIAN_BLUE_DEEP};
                border-radius: 40px;
                padding: 20px;
            """)

        logo_label.setAlignment(Qt.AlignmentFlag.AlignCenter)
        logo_label.setFixedSize(80, 80)

        logo_container = QWidget()
        logo_container_layout = QHBoxLayout(logo_container)
        logo_container_layout.setContentsMargins(0, 0, 0, 0)
        logo_container_layout.addStretch()
        logo_container_layout.addWidget(logo_label)
        logo_container_layout.addStretch()
        header_layout.addWidget(logo_container)

        # ROOT3POWER text
        company_label = QLabel("ROOT3POWER")
        company_label.setAlignment(Qt.AlignmentFlag.AlignCenter)
        company_label.setStyleSheet(f"""
            font-size: 20px;
            font-weight: bold;
            color: {OBSIDIAN_TEXT_ACCENT};
            letter-spacing: 2px;
            margin-top: 5px;
        """)
        header_layout.addWidget(company_label)

        # Transmittal Builder text
        app_label = QLabel("Transmittal Builder")
        app_label.setAlignment(Qt.AlignmentFlag.AlignCenter)
        app_label.setStyleSheet(f"""
            font-size: 14px;
            font-weight: normal;
            color: {OBSIDIAN_TEXT_SECONDARY};
            margin-bottom: 10px;
        """)
        header_layout.addWidget(app_label)

        layout.addWidget(header_container)
    
    def create_action_buttons(self, layout):
        """Create Preview and Generate buttons with consistent styling."""
        # Preview button (using default button styling for consistency)
        self.preview_btn = QPushButton(" Preview")
        self.preview_btn.setIcon(self.icon_manager.get_icon("file-text", 20, OBSIDIAN_TEXT_PRIMARY))
        self.preview_btn.setIconSize(QSize(20, 20))
        self.preview_btn.setMinimumHeight(40)
        self.preview_btn.setCursor(Qt.CursorShape.PointingHandCursor)
        self.preview_btn.clicked.connect(self.preview_transmittal)
        layout.addWidget(self.preview_btn)

        # Generate button (using default button styling for consistency)
        self.generate_btn = QPushButton(" Generate")
        self.generate_btn.setIcon(self.icon_manager.get_icon("rocket", 20, OBSIDIAN_TEXT_PRIMARY))
        self.generate_btn.setIconSize(QSize(20, 20))
        self.generate_btn.setMinimumHeight(40)
        self.generate_btn.setCursor(Qt.CursorShape.PointingHandCursor)
        self.generate_btn.clicked.connect(self.generate_transmittal)
        layout.addWidget(self.generate_btn)
    
    def create_status_indicator(self, layout):
        """Create status indicator label with modern styling."""
        status_container = QWidget()
        status_layout = QHBoxLayout(status_container)
        status_layout.setContentsMargins(0, 0, 0, 0)
        status_layout.setSpacing(8)

        # Status dot (using a styled label)
        self.status_dot = QLabel("●")
        self.status_dot.setStyleSheet(f"""
            color: {OBSIDIAN_SUCCESS};
            font-size: 16px;
        """)
        status_layout.addWidget(self.status_dot)

        # Status text
        self.status_label = QLabel("Ready")
        self.status_label.setStyleSheet(f"""
            color: {OBSIDIAN_TEXT_PRIMARY};
            font-weight: bold;
            font-size: 12px;
        """)
        status_layout.addWidget(self.status_label)
        status_layout.addStretch()

        layout.addWidget(status_container)
    
    def create_activity_log(self, layout):
        """Create activity log terminal."""
        # Label
        log_label = QLabel("Activity Log")
        log_label.setObjectName("sectionHeader")
        layout.addWidget(log_label)
        
        # Terminal text edit
        self.log_textbox = QTextEdit()
        self.log_textbox.setObjectName("terminal")
        self.log_textbox.setReadOnly(True)
        self.log_textbox.setFixedHeight(110)  # Optimized height
        layout.addWidget(self.log_textbox)
    
    def create_support_buttons(self, layout):
        """Create bug report, suggestion, and settings buttons with consistent styling."""
        # Buttons layout
        buttons_layout = QHBoxLayout()
        buttons_layout.setSpacing(10)

        # Bug report button (using default styling with text and white icon)
        self.bug_btn = QPushButton(" Report Bug")
        self.bug_btn.setIcon(self.icon_manager.get_icon("bug", 18, OBSIDIAN_TEXT_PRIMARY))
        self.bug_btn.setIconSize(QSize(18, 18))
        self.bug_btn.setMinimumHeight(40)
        self.bug_btn.setCursor(Qt.CursorShape.PointingHandCursor)
        self.bug_btn.clicked.connect(self.open_bug_report)
        buttons_layout.addWidget(self.bug_btn)

        # Suggestion button (using default styling with text and white icon)
        self.suggestion_btn = QPushButton(" Suggestion")
        self.suggestion_btn.setIcon(self.icon_manager.get_icon("lightbulb", 18, OBSIDIAN_TEXT_PRIMARY))
        self.suggestion_btn.setIconSize(QSize(18, 18))
        self.suggestion_btn.setMinimumHeight(40)
        self.suggestion_btn.setCursor(Qt.CursorShape.PointingHandCursor)
        self.suggestion_btn.clicked.connect(self.open_suggestion)
        buttons_layout.addWidget(self.suggestion_btn)

        # Settings button (using default styling with text and white icon)
        self.settings_btn = QPushButton(" Settings")
        self.settings_btn.setIcon(self.icon_manager.get_icon("settings", 18, OBSIDIAN_TEXT_PRIMARY))
        self.settings_btn.setIconSize(QSize(18, 18))
        self.settings_btn.setMinimumHeight(40)
        self.settings_btn.setCursor(Qt.CursorShape.PointingHandCursor)
        self.settings_btn.clicked.connect(self.open_settings)
        buttons_layout.addWidget(self.settings_btn)

        layout.addLayout(buttons_layout)
    
    def create_version_footer(self, layout):
        """Create version label at bottom right."""
        layout.addStretch()  # Push version to bottom

        version_label = QLabel(f"v{VERSION}")
        version_label.setStyleSheet(f"""
            font-size: 11px;
            color: {OBSIDIAN_TEXT_SECONDARY};
            padding: 10px 15px;
        """)
        version_label.setAlignment(Qt.AlignmentFlag.AlignRight | Qt.AlignmentFlag.AlignBottom)
        layout.addWidget(version_label)

    def create_transmittal_type_selector(self, layout):
        """Create transmittal type selector (Standard vs CID)."""
        section_container = self.create_section_frame("Transmittal Type", "layers")
        section_layout = QVBoxLayout(section_container)
        section_layout.setContentsMargins(20, 15, 20, 15)
        section_layout.setSpacing(15)

        # Radio button container
        radio_container = QWidget()
        radio_layout = QHBoxLayout(radio_container)
        radio_layout.setContentsMargins(0, 0, 0, 0)
        radio_layout.setSpacing(30)

        # Standard transmittal radio button
        self.standard_radio = QCheckBox("Standard Transmittal")
        self.standard_radio.setChecked(True)  # Default selection
        self.standard_radio.setStyleSheet(f"""
            QCheckBox {{
                color: {OBSIDIAN_TEXT_PRIMARY};
                font-size: 13px;
                font-weight: bold;
                spacing: 8px;
            }}
            QCheckBox::indicator {{
                width: 18px;
                height: 18px;
                border-radius: 9px;
                border: 2px solid {OBSIDIAN_BORDER};
                background: {OBSIDIAN_CARD_DEEP};
            }}
            QCheckBox::indicator:checked {{
                background: {ROOT3_BLUE};
                border: 2px solid {ROOT3_BLUE};
            }}
            QCheckBox::indicator:checked:after {{
                content: "";
                width: 8px;
                height: 8px;
                border-radius: 4px;
                background: white;
                position: absolute;
                top: 3px;
                left: 3px;
            }}
        """)
        self.standard_radio.toggled.connect(self.on_transmittal_type_changed)
        radio_layout.addWidget(self.standard_radio)

        # CID transmittal radio button
        self.cid_radio = QCheckBox("CID Transmittal")
        self.cid_radio.setStyleSheet(self.standard_radio.styleSheet())
        self.cid_radio.toggled.connect(self.on_transmittal_type_changed)
        radio_layout.addWidget(self.cid_radio)

        radio_layout.addStretch()
        section_layout.addWidget(radio_container)

        # Description text
        desc_label = QLabel("Standard: Uses PDF files with Excel index • CID: Uses .cid files with built-in document entry")
        desc_label.setStyleSheet(f"""
            color: {OBSIDIAN_TEXT_SECONDARY};
            font-size: 11px;
            font-style: italic;
        """)
        desc_label.setWordWrap(True)
        section_layout.addWidget(desc_label)

        layout.addWidget(section_container)

        # Store transmittal type
        self.transmittal_type = "standard"

    def on_transmittal_type_changed(self):
        """Handle transmittal type radio button changes."""
        # Ensure only one radio button is checked (manual radio button behavior)
        if self.sender() == self.standard_radio and self.standard_radio.isChecked():
            self.cid_radio.setChecked(False)
            self.transmittal_type = "standard"
            self.log_message("Switched to Standard Transmittal mode", "info")
        elif self.sender() == self.cid_radio and self.cid_radio.isChecked():
            self.standard_radio.setChecked(False)
            self.transmittal_type = "cid"
            self.log_message("Switched to CID Transmittal mode", "info")

        # Update UI visibility based on mode
        self.update_file_section_visibility()

    def update_file_section_visibility(self):
        """Update file section visibility based on transmittal type."""
        if hasattr(self, 'pdf_container') and hasattr(self, 'excel_container'):
            if self.transmittal_type == "standard":
                self.pdf_container.show()
                self.excel_container.show()
                if hasattr(self, 'cid_container'):
                    self.cid_container.hide()
                if hasattr(self, 'cid_document_table_container'):
                    self.cid_document_table_container.hide()
            else:  # CID mode
                self.pdf_container.hide()
                self.excel_container.hide()
                if hasattr(self, 'cid_container'):
                    self.cid_container.show()
                if hasattr(self, 'cid_document_table_container'):
                    self.cid_document_table_container.show()

    def create_main_content(self):
        """Create main content area with all form sections."""
        self.main_content_frame = QFrame()
        self.main_content_frame.setObjectName("mainContent")

        # Create scroll area for main content
        scroll_area = QScrollArea()
        scroll_area.setWidgetResizable(True)
        scroll_area.setFrameShape(QFrame.Shape.NoFrame)
        scroll_area.setHorizontalScrollBarPolicy(Qt.ScrollBarPolicy.ScrollBarAlwaysOff)
        self.main_scroll = scroll_area  # Store reference for validation focus

        # Create container widget for scrollable content
        scroll_widget = QWidget()
        scroll_layout = QVBoxLayout(scroll_widget)
        scroll_layout.setContentsMargins(25, 15, 25, 15)
        scroll_layout.setSpacing(20)

        # Add all sections
        self.create_transmittal_type_selector(scroll_layout)
        self.create_file_selection_section(scroll_layout)
        self.create_cid_document_table_section(scroll_layout)
        self.create_project_info_section(scroll_layout)
        self.create_from_info_section(scroll_layout)
        self.create_contacts_section(scroll_layout)
        self.create_transmittal_options_section(scroll_layout)

        # Add stretch at bottom
        scroll_layout.addStretch()

        # Set scroll widget
        scroll_area.setWidget(scroll_widget)

        # Add scroll area to main content frame
        main_layout = QVBoxLayout(self.main_content_frame)
        main_layout.setContentsMargins(0, 0, 0, 0)
        main_layout.addWidget(scroll_area)
    
    # ========================================================================
    # FORM SECTIONS
    # ========================================================================

    def create_file_selection_section(self, layout):
        """Create file selection section with New Session button."""
        section_container = self.create_section_frame("File Selection", "folder")

        # Add New Session button to the title area (matches Browse button style)
        title_container = section_container.findChild(QWidget)
        if title_container:
            title_layout = title_container.layout()
            if title_layout:
                # Create New Session button (same style as Browse buttons)
                new_session_btn = QPushButton("New Session")
                new_session_btn.setFixedWidth(110)  # Slightly wider than Browse (80)
                new_session_btn.setCursor(Qt.CursorShape.PointingHandCursor)
                # No custom stylesheet - will use the default button style from theme
                new_session_btn.clicked.connect(self.new_session)
                title_layout.addWidget(new_session_btn)

        section_layout = QVBoxLayout(section_container._content_widget)
        section_layout.setContentsMargins(15, 15, 15, 15)
        section_layout.setSpacing(12)

        # Template file (common to both modes)
        template_layout = QHBoxLayout()
        template_label = QLabel("Template File:")
        template_label.setFixedWidth(120)
        # Use DropLineEdit for drag & drop support
        self.template_entry = DropLineEdit(mode="file", patterns=["*.docx"])
        self.template_entry.setPlaceholderText("Browse or drag & drop template file...")
        template_btn = QPushButton("Browse")
        template_btn.setFixedWidth(80)
        template_btn.setCursor(Qt.CursorShape.PointingHandCursor)
        template_btn.clicked.connect(self.browse_template)
        template_layout.addWidget(template_label)
        template_layout.addWidget(self.template_entry, 1)
        template_layout.addWidget(template_btn)
        section_layout.addLayout(template_layout)

        # Standard mode inputs (PDF + Excel)
        self.pdf_container = QWidget()
        pdf_layout = QVBoxLayout(self.pdf_container)
        pdf_layout.setContentsMargins(0, 0, 0, 0)
        pdf_layout.setSpacing(12)

        # Index file
        index_layout = QHBoxLayout()
        index_label = QLabel("Index File:")
        index_label.setFixedWidth(120)
        # Use DropLineEdit for drag & drop support
        self.index_entry = DropLineEdit(mode="file", patterns=["*.xlsx", "*.xls"])
        self.index_entry.setPlaceholderText("Browse or drag & drop Excel file...")
        index_btn = QPushButton("Browse")
        index_btn.setFixedWidth(80)
        index_btn.setCursor(Qt.CursorShape.PointingHandCursor)
        index_btn.clicked.connect(self.browse_index)
        index_layout.addWidget(index_label)
        index_layout.addWidget(self.index_entry, 1)
        index_layout.addWidget(index_btn)
        pdf_layout.addLayout(index_layout)

        # PDF folder
        pdf_folder_layout = QHBoxLayout()
        pdf_label = QLabel("PDF Folder:")
        pdf_label.setFixedWidth(120)
        # Use DropLineEdit for drag & drop support (directory mode)
        self.output_entry = DropLineEdit(mode="dir")
        self.output_entry.setPlaceholderText("Browse or drag & drop PDF folder...")
        output_btn = QPushButton("Browse")
        output_btn.setFixedWidth(80)
        output_btn.setCursor(Qt.CursorShape.PointingHandCursor)
        output_btn.clicked.connect(self.browse_output)
        pdf_folder_layout.addWidget(pdf_label)
        pdf_folder_layout.addWidget(self.output_entry, 1)
        pdf_folder_layout.addWidget(output_btn)
        pdf_layout.addLayout(pdf_folder_layout)

        self.excel_container = self.pdf_container  # Excel is part of standard mode
        section_layout.addWidget(self.pdf_container)

        # CID mode inputs
        self.cid_container = QWidget()
        cid_layout = QVBoxLayout(self.cid_container)
        cid_layout.setContentsMargins(0, 0, 0, 0)
        cid_layout.setSpacing(12)

        # CID folder
        cid_folder_layout = QHBoxLayout()
        cid_label = QLabel("CID Folder:")
        cid_label.setFixedWidth(120)
        # Use DropLineEdit for drag & drop support (directory mode)
        self.cid_entry = DropLineEdit(mode="dir")
        self.cid_entry.setPlaceholderText("Browse or drag & drop CID folder...")
        cid_btn = QPushButton("Browse")
        cid_btn.setFixedWidth(80)
        cid_btn.setCursor(Qt.CursorShape.PointingHandCursor)
        cid_btn.clicked.connect(self.browse_cid_folder)
        cid_folder_layout.addWidget(cid_label)
        cid_folder_layout.addWidget(self.cid_entry, 1)
        cid_folder_layout.addWidget(cid_btn)
        cid_layout.addLayout(cid_folder_layout)

        section_layout.addWidget(self.cid_container)
        self.cid_container.hide()  # Hidden by default

        layout.addWidget(section_container)

    def create_cid_document_table_section(self, layout):
        """Create CID document entry table section."""
        self.cid_document_table_container = self.create_section_frame("CID Document Index", "file-text")
        section_layout = QVBoxLayout(self.cid_document_table_container)
        section_layout.setContentsMargins(20, 15, 20, 15)
        section_layout.setSpacing(15)

        # Instructions
        instructions = QLabel("CID files will be automatically scanned. Enter document descriptions and revisions below:")
        instructions.setStyleSheet(f"""
            color: {OBSIDIAN_TEXT_SECONDARY};
            font-size: 12px;
            margin-bottom: 10px;
        """)
        instructions.setWordWrap(True)
        section_layout.addWidget(instructions)

        # Scroll area for the table
        scroll_area = QScrollArea()
        scroll_area.setWidgetResizable(True)
        scroll_area.setFrameShape(QFrame.Shape.NoFrame)
        scroll_area.setMaximumHeight(300)
        scroll_area.setStyleSheet(f"""
            QScrollArea {{
                background: {OBSIDIAN_CARD_DEEP};
                border: 1px solid {OBSIDIAN_BORDER};
                border-radius: 6px;
            }}
        """)

        # Table widget
        self.cid_table_widget = QWidget()
        self.cid_table_layout = QVBoxLayout(self.cid_table_widget)
        self.cid_table_layout.setContentsMargins(10, 10, 10, 10)
        self.cid_table_layout.setSpacing(8)

        # Table header
        header_widget = QWidget()
        header_layout = QHBoxLayout(header_widget)
        header_layout.setContentsMargins(5, 5, 5, 5)
        header_layout.setSpacing(10)

        # Header labels
        file_header = QLabel("CID File")
        file_header.setFixedWidth(200)
        file_header.setStyleSheet(f"""
            font-weight: bold;
            color: {OBSIDIAN_TEXT_ACCENT};
            padding: 5px;
        """)
        header_layout.addWidget(file_header)

        desc_header = QLabel("Document Description")
        desc_header.setStyleSheet(f"""
            font-weight: bold;
            color: {OBSIDIAN_TEXT_ACCENT};
            padding: 5px;
        """)
        header_layout.addWidget(desc_header, 1)

        rev_header = QLabel("Revision")
        rev_header.setFixedWidth(100)
        rev_header.setStyleSheet(f"""
            font-weight: bold;
            color: {OBSIDIAN_TEXT_ACCENT};
            padding: 5px;
        """)
        header_layout.addWidget(rev_header)

        self.cid_table_layout.addWidget(header_widget)

        # Container for CID document rows
        self.cid_document_rows = []

        scroll_area.setWidget(self.cid_table_widget)
        section_layout.addWidget(scroll_area)

        # Buttons
        button_layout = QHBoxLayout()

        # Scan CID folder button
        scan_btn = QPushButton(" Scan CID Folder")
        scan_btn.setIcon(self.icon_manager.get_icon("refresh-cw", 18, OBSIDIAN_TEXT_PRIMARY))
        scan_btn.setIconSize(QSize(18, 18))
        scan_btn.setCursor(Qt.CursorShape.PointingHandCursor)
        scan_btn.clicked.connect(self.scan_cid_folder)
        button_layout.addWidget(scan_btn)

        # Auto-rename button
        rename_btn = QPushButton(" Auto-Rename CIDs")
        rename_btn.setIcon(self.icon_manager.get_icon("edit", 18, OBSIDIAN_TEXT_PRIMARY))
        rename_btn.setIconSize(QSize(18, 18))
        rename_btn.setCursor(Qt.CursorShape.PointingHandCursor)
        rename_btn.clicked.connect(self.auto_rename_cid_files)
        button_layout.addWidget(rename_btn)

        button_layout.addStretch()
        section_layout.addLayout(button_layout)

        layout.addWidget(self.cid_document_table_container)
        self.cid_document_table_container.hide()  # Hidden by default

    def browse_cid_folder(self):
        """Browse for CID folder."""
        folder = QFileDialog.getExistingDirectory(self, "Select CID Folder")
        if folder:
            self.cid_entry.setText(folder)
            self.log_message(f"CID folder selected: {folder}", "success")
            # Auto-scan when folder is selected
            self.scan_cid_folder()

    def scan_cid_folder(self):
        """Scan CID folder and populate document table."""
        cid_folder = self.cid_entry.text().strip()
        if not cid_folder or not os.path.exists(cid_folder):
            self.log_message("Please select a valid CID folder first", "error")
            return

        # Clear existing rows
        self.clear_cid_document_rows()

        # Find all .cid files
        cid_files = []
        try:
            for file in os.listdir(cid_folder):
                if file.lower().endswith('.cid'):
                    cid_files.append(file)
        except Exception as e:
            self.log_message(f"Error scanning CID folder: {e}", "error")
            return

        if not cid_files:
            self.log_message("No .cid files found in selected folder", "warning")
            return

        # Sort files for consistent ordering
        cid_files.sort()

        # Create table rows for each CID file
        for cid_file in cid_files:
            self.add_cid_document_row(cid_file)

        self.log_message(f"Found {len(cid_files)} CID files", "success")

    def clear_cid_document_rows(self):
        """Clear all CID document rows."""
        for row_data in self.cid_document_rows:
            row_data['widget'].setParent(None)
        self.cid_document_rows.clear()

    def add_cid_document_row(self, cid_filename):
        """Add a CID document row to the table."""
        row_widget = QWidget()
        row_layout = QHBoxLayout(row_widget)
        row_layout.setContentsMargins(5, 5, 5, 5)
        row_layout.setSpacing(10)

        # CID filename (read-only)
        filename_label = QLabel(cid_filename)
        filename_label.setFixedWidth(200)
        filename_label.setStyleSheet(f"""
            color: {OBSIDIAN_TEXT_PRIMARY};
            background: {OBSIDIAN_CARD_MID};
            border: 1px solid {OBSIDIAN_BORDER};
            border-radius: 4px;
            padding: 5px;
            font-family: monospace;
        """)
        row_layout.addWidget(filename_label)

        # Document description (editable)
        desc_entry = QLineEdit()
        desc_entry.setPlaceholderText("Enter document description...")
        desc_entry.setStyleSheet(f"""
            QLineEdit {{
                background: {OBSIDIAN_CARD_DEEP};
                color: {OBSIDIAN_TEXT_PRIMARY};
                border: 1px solid {OBSIDIAN_BORDER};
                border-radius: 4px;
                padding: 5px;
            }}
            QLineEdit:focus {{
                border: 2px solid {OBSIDIAN_BLUE_GLOW};
            }}
        """)

        # Auto-generate description from filename
        auto_description = self.parse_cid_filename(cid_filename)
        if auto_description:
            desc_entry.setText(auto_description)

        row_layout.addWidget(desc_entry, 1)

        # Revision (combobox)
        rev_combo = QComboBox()
        rev_combo.setFixedWidth(100)
        rev_combo.addItems(["-", "0", "1", "2", "3", "4", "5", "A", "B", "C", "D", "E", "IFA", "IFC"])
        rev_combo.setCurrentText("-")
        rev_combo.setStyleSheet(f"""
            QComboBox {{
                background: {OBSIDIAN_CARD_DEEP};
                color: {OBSIDIAN_TEXT_PRIMARY};
                border: 1px solid {OBSIDIAN_BORDER};
                border-radius: 4px;
                padding: 5px;
            }}
            QComboBox:focus {{
                border: 2px solid {OBSIDIAN_BLUE_GLOW};
            }}
            QComboBox::drop-down {{
                border: none;
                width: 20px;
            }}
            QComboBox::down-arrow {{
                image: none;
                border: none;
            }}
        """)
        row_layout.addWidget(rev_combo)

        # Store row data
        row_data = {
            'widget': row_widget,
            'filename': cid_filename,
            'filename_label': filename_label,
            'description': desc_entry,
            'revision': rev_combo
        }
        self.cid_document_rows.append(row_data)

        # Add to layout
        self.cid_table_layout.addWidget(row_widget)

    def parse_cid_filename(self, filename):
        """Parse CID filename to generate document description."""
        try:
            # Remove .cid extension
            name = filename.replace('.cid', '')

            # Remove R3P_YYYYMMDD suffix if present
            import re
            name = re.sub(r'_R3P_\d{8}$', '', name)

            # Parse pattern: 850_LV-SWG-SG61-2A_M1
            parts = name.split('_')
            if len(parts) >= 3:
                relay_type = parts[0]
                relay_name = parts[1]
                feeds_protects = parts[2]

                return f"Relay {relay_type} ({relay_name}) protecting {feeds_protects}"
            elif len(parts) == 2:
                relay_type = parts[0]
                relay_name = parts[1]
                return f"Relay {relay_type} ({relay_name})"
            else:
                return f"Relay {name}"

        except Exception:
            # Fallback to filename without extension
            return filename.replace('.cid', '')

    def auto_rename_cid_files(self):
        """Auto-rename CID files to include R3P_YYYYMMDD suffix."""
        cid_folder = self.cid_entry.text().strip()
        if not cid_folder or not os.path.exists(cid_folder):
            self.log_message("Please select a valid CID folder first", "error")
            return

        # Get today's date in YYYYMMDD format
        from datetime import datetime
        today = datetime.now().strftime("%Y%m%d")
        suffix = f"_R3P_{today}"

        renamed_count = 0
        skipped_count = 0

        try:
            for file in os.listdir(cid_folder):
                if file.lower().endswith('.cid'):
                    # Check if already has R3P suffix
                    if '_R3P_' in file:
                        skipped_count += 1
                        continue

                    # Generate new filename
                    name_without_ext = file.replace('.cid', '')
                    new_filename = f"{name_without_ext}{suffix}.cid"

                    old_path = os.path.join(cid_folder, file)
                    new_path = os.path.join(cid_folder, new_filename)

                    # Rename the file
                    os.rename(old_path, new_path)
                    renamed_count += 1

            # Log results
            if renamed_count > 0:
                self.log_message(f"Renamed {renamed_count} CID files with R3P_{today} suffix", "success")
            if skipped_count > 0:
                self.log_message(f"Skipped {skipped_count} files (already have R3P suffix)", "info")

            # Refresh the table if files were renamed
            if renamed_count > 0:
                self.scan_cid_folder()

        except Exception as e:
            self.log_message(f"Error renaming CID files: {e}", "error")

    def collect_cid_index_data(self):
        """Collect CID document index data from the table."""
        index_data = []
        for row_data in self.cid_document_rows:
            filename = row_data['filename']
            description = row_data['description'].text().strip()
            revision = row_data['revision'].currentText()

            if description:  # Only include rows with descriptions
                index_data.append({
                    'filename': filename,
                    'description': description,
                    'revision': revision if revision != "-" else ""
                })

        return index_data

    def create_project_info_section(self, layout):
        """Create project information section."""
        section_container = self.create_section_frame("Project Information", "clipboard")
        section_layout = QGridLayout(section_container._content_widget)
        section_layout.setContentsMargins(15, 15, 15, 15)
        section_layout.setSpacing(12)

        # Project name
        section_layout.addWidget(QLabel("Project Name:"), 0, 0)
        self.project_name_entry = QLineEdit()
        self.project_name_entry.setPlaceholderText("<Client> - <Site Name>")
        self.project_name_entry.textChanged.connect(self.trigger_auto_save)
        section_layout.addWidget(self.project_name_entry, 0, 1)

        # Project number
        section_layout.addWidget(QLabel("Project Number:"), 0, 2)
        self.project_number_entry = QLineEdit()
        self.project_number_entry.setPlaceholderText("R3P-<PRJ#>")
        self.project_number_entry.textChanged.connect(self.trigger_auto_save)
        section_layout.addWidget(self.project_number_entry, 0, 3)

        # Date
        section_layout.addWidget(QLabel("Date:"), 1, 0)
        self.date_entry = QLineEdit()
        self.date_entry.setText(datetime.now().strftime("%m/%d/%Y"))
        self.date_entry.textChanged.connect(self.trigger_auto_save)
        section_layout.addWidget(self.date_entry, 1, 1)

        # Transmittal (NEW!)
        section_layout.addWidget(QLabel("Transmittal:"), 1, 2)
        self.transmittal_entry = QLineEdit()
        self.transmittal_entry.setPlaceholderText("XMTL-<###>")
        self.transmittal_entry.textChanged.connect(self.trigger_auto_save)
        section_layout.addWidget(self.transmittal_entry, 1, 3)

        # Description
        section_layout.addWidget(QLabel("Description:"), 2, 0, Qt.AlignmentFlag.AlignTop)
        self.description_text = QTextEdit()
        self.description_text.setPlaceholderText("<PROJECT DESCRIPTION>")
        self.description_text.setMaximumHeight(80)
        self.description_text.textChanged.connect(self.trigger_auto_save)
        section_layout.addWidget(self.description_text, 2, 1, 1, 3)

        layout.addWidget(section_container)

    def create_from_info_section(self, layout):
        """Create from information section."""
        section_container = self.create_section_frame("From Information", "user")
        section_layout = QGridLayout(section_container._content_widget)
        section_layout.setContentsMargins(15, 15, 15, 15)
        section_layout.setSpacing(12)

        # Hidden from_name field (populated by PE selection)
        self.from_name_entry = QLineEdit()
        self.from_name_entry.setVisible(False)

        # PE dropdown (no scroll) - moved to first position
        section_layout.addWidget(QLabel("PE:"), 0, 0)
        self.pe_combo = NoScrollComboBox()
        # Load PE names from config and sort alphabetically
        pe_names = config_manager.get_pe_names() if config_manager else []
        # Remove "None" if present and sort
        pe_names = [name for name in pe_names if name != "None"]
        pe_names.sort()
        # Add blank option at the beginning
        self.pe_combo.addItem("")  # Blank option
        self.pe_combo.addItems(pe_names)
        self.pe_combo.setCurrentIndex(0)  # Start with blank
        self.pe_combo.currentTextChanged.connect(self.on_pe_changed)
        section_layout.addWidget(self.pe_combo, 0, 1)

        # Title
        section_layout.addWidget(QLabel("Title:"), 0, 2)
        self.from_title_entry = QLineEdit()
        self.from_title_entry.textChanged.connect(self.trigger_auto_save)
        section_layout.addWidget(self.from_title_entry, 0, 3)

        # Email
        section_layout.addWidget(QLabel("Email:"), 1, 0)
        self.from_email_entry = QLineEdit()
        self.from_email_entry.textChanged.connect(self.trigger_auto_save)
        section_layout.addWidget(self.from_email_entry, 1, 1)

        # Phone
        section_layout.addWidget(QLabel("Phone:"), 1, 2)
        self.from_phone_entry = QLineEdit()
        self.from_phone_entry.textChanged.connect(self.trigger_auto_save)
        section_layout.addWidget(self.from_phone_entry, 1, 3)

        # Firm number by state (no scroll)
        section_layout.addWidget(QLabel("Firm Number:"), 2, 0)
        self.firm_combo = NoScrollComboBox()
        # Load firm numbers from config and sort alphabetically
        firm_numbers = config_manager.get_firm_numbers() if config_manager else []
        # Remove "None" if present and sort
        firm_numbers = [num for num in firm_numbers if num != "None"]
        firm_numbers.sort()
        # Add blank option at the beginning
        self.firm_combo.addItem("")  # Blank option
        self.firm_combo.addItems(firm_numbers)
        self.firm_combo.setCurrentIndex(0)  # Start with blank
        self.firm_combo.currentTextChanged.connect(self.trigger_auto_save)
        section_layout.addWidget(self.firm_combo, 2, 1, 1, 3)  # Span across remaining columns

        layout.addWidget(section_container)

    def create_contacts_section(self, layout):
        """Create contacts section with dynamic contact list."""
        section_container = self.create_section_frame("To - Contacts", "users")
        section_layout = QVBoxLayout(section_container._content_widget)
        section_layout.setContentsMargins(15, 15, 15, 15)
        section_layout.setSpacing(12)

        # Container for contact rows
        self.contacts_container = QWidget()
        self.contacts_layout = QVBoxLayout(self.contacts_container)
        self.contacts_layout.setSpacing(8)
        section_layout.addWidget(self.contacts_container)

        # Add contact button
        add_btn = QPushButton(" Add Contact")
        add_btn.setIcon(self.icon_manager.get_icon("plus", 18, OBSIDIAN_TEXT_PRIMARY))
        add_btn.setIconSize(QSize(18, 18))
        add_btn.setCursor(Qt.CursorShape.PointingHandCursor)
        add_btn.clicked.connect(self.add_contact_row)
        section_layout.addWidget(add_btn)

        # Initialize with 1 contact (suppress logging during init)
        self.contact_rows = []
        self._suppress_contact_log = True
        self.add_contact_row()
        self._suppress_contact_log = False
        # No log message - contacts are added silently during startup

        layout.addWidget(section_container)

    def create_transmittal_options_section(self, layout):
        """Create transmittal options section."""
        section_container = self.create_section_frame("Transmittal Options", "settings")
        section_layout = QGridLayout(section_container._content_widget)
        section_layout.setContentsMargins(15, 15, 15, 15)
        section_layout.setSpacing(15)

        # Create option groups
        col = 0

        # Transmitted
        transmitted_group = self.create_option_group("Transmitted", [
            ("PDF", "trans_pdf"),
            ("CAD", "trans_cad"),
            ("Originals", "trans_originals")
        ], "send")
        section_layout.addWidget(transmitted_group, 0, col)
        col += 1

        # Sent Via
        sent_via_group = self.create_option_group("Sent Via", [
            ("Email", "via_email"),
            ("FTP", "via_ftp")
        ], "mail")
        section_layout.addWidget(sent_via_group, 0, col)
        col += 1

        # Client Issue (reordered as requested)
        client_issue_group = self.create_option_group("Client Issue", [
            ("For Bid", "ci_bid"),
            ("For Preliminary", "ci_preliminary"),
            ("For Approval", "ci_approval"),
            ("For Construction", "ci_construction"),
            ("For As-Built", "ci_asbuilt"),
            ("For Information Only", "ci_info"),
            ("For Reference", "ci_reference")
        ], "clipboard")
        section_layout.addWidget(client_issue_group, 0, col)
        col += 1

        # Vendor Return
        vendor_return_group = self.create_option_group("Vendor Return", [
            ("Approved", "vr_approved"),
            ("Approved as Noted", "vr_approved_noted"),
            ("Rejected", "vr_rejected")
        ], "corner-down-left")
        section_layout.addWidget(vendor_return_group, 0, col)

        layout.addWidget(section_container)

    # ========================================================================
    # HELPER METHODS FOR UI CREATION
    # ========================================================================

    def add_drop_shadow(self, widget: QWidget, blur_radius: int = 20,
                       offset_x: int = 0, offset_y: int = 4,
                       color: QColor = None):
        """Add a drop shadow effect to a widget."""
        if color is None:
            color = QColor(0, 0, 0, 80)  # Semi-transparent black

        shadow = QGraphicsDropShadowEffect()
        shadow.setBlurRadius(blur_radius)
        shadow.setColor(color)
        shadow.setOffset(offset_x, offset_y)
        widget.setGraphicsEffect(shadow)

    def add_glow_effect(self, widget: QWidget, color: QColor = None):
        """Add a glow effect to a widget."""
        if color is None:
            color = QColor(OBSIDIAN_BLUE_GLOW)

        glow = QGraphicsDropShadowEffect()
        glow.setBlurRadius(15)
        glow.setColor(color)
        glow.setOffset(0, 0)  # No offset for glow
        widget.setGraphicsEffect(glow)

    def create_section_frame(self, title: str, icon_name: str = None) -> QWidget:
        """Create a styled section frame with title and optional icon."""
        # Container widget
        container = QWidget()
        container_layout = QVBoxLayout(container)
        container_layout.setContentsMargins(0, 0, 0, 0)
        container_layout.setSpacing(0)

        # Card frame
        card_frame = QFrame()
        card_frame.setObjectName("glassCard")
        card_layout = QVBoxLayout(card_frame)
        card_layout.setContentsMargins(0, 0, 0, 0)
        card_layout.setSpacing(0)

        # Add drop shadow to card
        self.add_drop_shadow(card_frame, blur_radius=25, offset_y=6,
                            color=QColor(0, 0, 0, 100))

        # Title container with icon
        title_container = QWidget()
        title_layout = QHBoxLayout(title_container)
        title_layout.setContentsMargins(15, 12, 15, 12)
        title_layout.setSpacing(10)

        # Add icon if specified
        if icon_name:
            icon_label = QLabel()
            icon_pixmap = self.icon_manager.get_pixmap(icon_name, 20, OBSIDIAN_TEXT_ACCENT)
            icon_label.setPixmap(icon_pixmap)
            icon_label.setFixedSize(20, 20)
            title_layout.addWidget(icon_label)

        # Title label
        title_label = QLabel(title)
        title_label.setObjectName("sectionHeader")
        title_layout.addWidget(title_label)
        title_layout.addStretch()

        # Style the title container
        title_container.setStyleSheet(f"""
            background: qlineargradient(
                x1:0, y1:0, x2:1, y2:0,
                stop:0 {OBSIDIAN_CARD_DEEP},
                stop:1 {OBSIDIAN_CARD_MID}
            );
            border-bottom: 1px solid {OBSIDIAN_BORDER};
        """)

        card_layout.addWidget(title_container)

        # Content widget (where section content goes)
        content_widget = QWidget()
        card_layout.addWidget(content_widget)

        container_layout.addWidget(card_frame)

        # Return both container and content widget
        container._content_widget = content_widget
        return container

    def create_option_group(self, title: str, options: list, icon_name: str = None) -> QFrame:
        """Create an option group with checkboxes and optional icon."""
        group_frame = QFrame()
        group_layout = QVBoxLayout(group_frame)
        group_layout.setContentsMargins(10, 10, 10, 10)
        group_layout.setSpacing(8)

        # Title with icon
        title_container = QWidget()
        title_layout = QHBoxLayout(title_container)
        title_layout.setContentsMargins(0, 0, 0, 5)
        title_layout.setSpacing(6)

        # Add icon if specified
        if icon_name:
            icon_label = QLabel()
            icon_pixmap = self.icon_manager.get_pixmap(icon_name, 16, OBSIDIAN_TEXT_ACCENT)
            icon_label.setPixmap(icon_pixmap)
            icon_label.setFixedSize(16, 16)
            title_layout.addWidget(icon_label)

        # Title label
        title_label = QLabel(title)
        title_label.setObjectName("sectionHeader")
        title_label.setStyleSheet(f"font-size: 12px;")
        title_layout.addWidget(title_label)
        title_layout.addStretch()

        group_layout.addWidget(title_container)

        # Checkboxes
        for label, attr_name in options:
            checkbox = QCheckBox(label)
            checkbox.stateChanged.connect(self.trigger_auto_save)
            setattr(self, f"chk_{attr_name}", checkbox)
            group_layout.addWidget(checkbox)

        group_layout.addStretch()
        return group_frame

    def add_contact_row(self):
        """Add a new contact row."""
        contact_widget = QWidget()
        contact_layout = QHBoxLayout(contact_widget)
        contact_layout.setContentsMargins(0, 0, 0, 0)
        contact_layout.setSpacing(8)

        # Name
        name_entry = QLineEdit()
        name_entry.setPlaceholderText("Name")
        name_entry.textChanged.connect(self.trigger_auto_save)
        name_entry.textChanged.connect(self._on_contact_field_changed)
        contact_layout.addWidget(name_entry, 2)

        # Company
        company_entry = QLineEdit()
        company_entry.setPlaceholderText("Company")
        company_entry.textChanged.connect(self.trigger_auto_save)
        company_entry.textChanged.connect(self._on_contact_field_changed)
        contact_layout.addWidget(company_entry, 2)

        # Email
        email_entry = QLineEdit()
        email_entry.setPlaceholderText("Email")
        email_entry.textChanged.connect(self.trigger_auto_save)
        email_entry.textChanged.connect(self._on_contact_field_changed)
        contact_layout.addWidget(email_entry, 2)

        # Phone
        phone_entry = QLineEdit()
        phone_entry.setPlaceholderText("Phone")
        phone_entry.textChanged.connect(self.trigger_auto_save)
        phone_entry.textChanged.connect(self._on_contact_field_changed)
        contact_layout.addWidget(phone_entry, 1)

        # Remove button
        remove_btn = QPushButton()
        remove_btn.setIcon(self.icon_manager.get_icon("x", 16, OBSIDIAN_TEXT_PRIMARY))
        remove_btn.setIconSize(QSize(16, 16))
        remove_btn.setFixedSize(30, 30)
        remove_btn.setCursor(Qt.CursorShape.PointingHandCursor)
        remove_btn.setStyleSheet(f"""
            QPushButton {{
                background-color: {OBSIDIAN_ERROR};
                border-radius: 4px;
            }}
            QPushButton:hover {{
                background-color: #FF6B7A;
            }}
        """)
        remove_btn.clicked.connect(lambda: self.remove_contact_row(contact_widget))
        contact_layout.addWidget(remove_btn)

        # Store references
        contact_data = {
            'widget': contact_widget,
            'name': name_entry,
            'company': company_entry,
            'email': email_entry,
            'phone': phone_entry
        }
        self.contact_rows.append(contact_data)

        # Add to layout
        self.contacts_layout.addWidget(contact_widget)

        # Update window size
        self.update_window_size()

        # Only log if not suppressed (during init)
        if not getattr(self, '_suppress_contact_log', False):
            self.log_message(f"Contact added (Total: {len(self.contact_rows)})", "info")

    def remove_contact_row(self, widget):
        """Remove a contact row."""
        from utils.validation import _mark_valid

        for i, contact_data in enumerate(self.contact_rows):
            if contact_data['widget'] == widget:
                # Clear red borders before removing
                _mark_valid(contact_data['name'])
                _mark_valid(contact_data['company'])
                _mark_valid(contact_data['email'])
                _mark_valid(contact_data['phone'])

                widget.deleteLater()
                self.contact_rows.pop(i)
                self.update_window_size()

                # Only log if not suppressed (during init or new session)
                if not getattr(self, '_suppress_contact_log', False):
                    self.log_message(f"Contact removed (Total: {len(self.contact_rows)})", "info")

                # Update validation after removal
                if self._show_validation:
                    self._update_field_borders()

                break

    def update_window_size(self):
        """Update window size based on contact count and screen size."""
        # Get current screen info
        screen = QApplication.primaryScreen()
        screen_geometry = screen.availableGeometry()

        # Base height depends on screen size (updated to match new sizing)
        if screen_geometry.width() <= 1366:
            base_height = 750
        elif screen_geometry.width() <= 1920:
            base_height = 850
        else:
            base_height = 950

        contact_count = len(self.contact_rows)

        if contact_count > 3:
            extra_height = (contact_count - 3) * 60
            new_height = min(base_height + extra_height, int(screen_geometry.height() * 0.9))
            self.resize(self.width(), new_height)

    # ========================================================================
    # DATA COLLECTION METHODS
    # ========================================================================

    def collect_fields(self) -> dict:
        """Collect all form fields into a dictionary."""
        return {
            "date": self.date_entry.text(),
            "job_num": self.project_number_entry.text(),
            "transmittal_num": self.transmittal_entry.text() if hasattr(self, 'transmittal_entry') else "",
            "client": self.project_name_entry.text(),
            "project_desc": self.description_text.toPlainText(),
            "from_name": self.from_name_entry.text(),
            "from_title": self.from_title_entry.text(),
            "from_email": self.from_email_entry.text(),
            "from_phone": self.from_phone_entry.text(),
            "firm": self.firm_combo.currentText(),
        }

    def collect_checks(self) -> dict:
        """Collect all checkbox states into a dictionary."""
        return {
            # Transmitted
            "trans_pdf": self.chk_trans_pdf.isChecked(),
            "trans_cad": self.chk_trans_cad.isChecked(),
            "trans_originals": self.chk_trans_originals.isChecked(),

            # Sent Via
            "via_email": self.chk_via_email.isChecked(),
            "via_ftp": self.chk_via_ftp.isChecked(),

            # Client Issue (PyQt6 version)
            "ci_approval": self.chk_ci_approval.isChecked(),
            "ci_bid": self.chk_ci_bid.isChecked(),
            "ci_construction": self.chk_ci_construction.isChecked(),
            "ci_asbuilt": self.chk_ci_asbuilt.isChecked(),
            "ci_reference": self.chk_ci_reference.isChecked(),
            "ci_preliminary": self.chk_ci_preliminary.isChecked(),
            "ci_info": self.chk_ci_info.isChecked(),

            # Client Issue (old GUI compatibility - provide defaults)
            "ci_fab": False,
            "ci_const": self.chk_ci_construction.isChecked(),  # Map to construction
            "ci_record": False,
            "ci_ref": self.chk_ci_reference.isChecked(),  # Map to reference

            # Vendor Return
            "vr_approved": self.chk_vr_approved.isChecked(),
            "vr_approved_noted": self.chk_vr_approved_noted.isChecked(),
            "vr_rejected": self.chk_vr_rejected.isChecked(),
        }

    def collect_contacts(self) -> List[Dict[str, str]]:
        """Collect all contact information into a list."""
        contacts = []
        for contact_data in self.contact_rows:
            name = contact_data['name'].text().strip()
            company = contact_data['company'].text().strip()
            email = contact_data['email'].text().strip()
            phone = contact_data['phone'].text().strip()

            # Only add if ALL required fields are filled (name, company, email, phone)
            if name and company and email and phone:
                contacts.append({
                    "name": name,
                    "company": company,
                    "email": email,
                    "phone": phone
                })
        return contacts

    def validate_required_fields(self) -> Tuple[bool, str]:
        """Validate that all required fields are filled. Returns (is_valid, error_message)."""
        # Check template file
        if not self.template_entry.text().strip():
            return False, "Please select a template file."

        if not os.path.isfile(self.template_entry.text()):
            return False, "Template file does not exist."

        # Check index file
        if not self.index_entry.text().strip():
            return False, "Please select a drawing index file."

        if not os.path.isfile(self.index_entry.text()):
            return False, "Drawing index file does not exist."

        # Check output folder
        if not self.output_entry.text().strip():
            return False, "Please select an output folder."

        # Check project info
        if not self.project_name_entry.text().strip():
            return False, "Please enter a project name."

        if not self.project_number_entry.text().strip():
            return False, "Please enter a project number."

        # Check from info (PE selection)
        if not self.from_name_entry.text().strip():
            return False, "Please select a PE in the From Information section."

        # Check contacts
        contacts = self.collect_contacts()
        if not contacts:
            return False, "Please add at least one contact."

        return True, ""

    # ========================================================================
    # UTILITY METHODS
    # ========================================================================

    def update_status(self, text: str, color: str = None):
        """Update the status indicator."""
        if color is None:
            color = OBSIDIAN_SUCCESS

        self.status_label.setText(text)
        self.status_dot.setStyleSheet(f"""
            color: {color};
            font-size: 16px;
        """)

    def log_message(self, message: str, level: str = "info"):
        """Add message to activity log."""
        timestamp = datetime.now().strftime("%H:%M:%S")

        # Icon and color based on level
        if level == "success":
            color = OBSIDIAN_SUCCESS
            icon = "●"  # Using simple circle for now (SVG icons in text are complex)
        elif level == "error":
            color = OBSIDIAN_ERROR
            icon = "●"
        elif level == "warning":
            color = OBSIDIAN_WARNING
            icon = "●"
        else:
            color = OBSIDIAN_TEXT_ACCENT
            icon = "●"

        # Format message with HTML
        formatted_msg = f'<span style="color: {OBSIDIAN_TEXT_SECONDARY}">{timestamp}</span> <span style="color: {color}">{icon}</span> {message}'

        # Append to log
        self.log_textbox.append(formatted_msg)

        # Auto-scroll to bottom
        scrollbar = self.log_textbox.verticalScrollBar()
        scrollbar.setValue(scrollbar.maximum())
    
    def setup_auto_save(self):
        """Setup auto-save timer."""
        # Get auto-save interval from config (in seconds)
        interval_seconds = config_manager.get_auto_save_interval() if config_manager else 120
        interval_ms = interval_seconds * 1000

        self.auto_save_timer = QTimer()
        self.auto_save_timer.timeout.connect(self.auto_save)
        self.auto_save_timer.start(interval_ms)
        self.auto_save_interval_ms = interval_ms  # Store for trigger_auto_save

    def trigger_auto_save(self):
        """Trigger auto-save (debounced)."""
        # Reset timer to debounce rapid changes
        self.auto_save_timer.stop()
        self.auto_save_timer.start(self.auto_save_interval_ms)

    def setup_validation(self):
        """Setup live validation for form fields (pristine mode)."""

        def update_validation():
            """Update validation state and Generate button."""
            # Check all fields in real-time
            all_valid = (
                self._is_valid_email() and
                self._is_valid_project_name() and
                self._is_valid_template() and
                self._is_valid_index() and
                self._is_valid_output()
            )
            self.generate_btn.setEnabled(all_valid)
            if not all_valid:
                self.generate_btn.setToolTip("Please complete required fields")
            else:
                self.generate_btn.setToolTip("")

            # If validation is visible, update red borders
            if self._show_validation:
                self._update_field_borders()

        # Validate From Email
        run1 = attach_line_validator(
            self.from_email_entry,
            lambda text: is_valid_email(text) and bool(text.strip()),
            "Valid email address required",
            update_validation,
            is_visual=lambda: self._show_validation,
            debounce_ms=300
        )
        self._validator_runs.append(run1)

        # Validate Project Name
        run2 = attach_line_validator(
            self.project_name_entry,
            lambda text: bool(text.strip()),
            "Project name required",
            update_validation,
            is_visual=lambda: self._show_validation,
            debounce_ms=300
        )
        self._validator_runs.append(run2)

        # Validate Project Number
        run3 = attach_line_validator(
            self.project_number_entry,
            lambda text: bool(text.strip()),
            "Project number required",
            update_validation,
            is_visual=lambda: self._show_validation,
            debounce_ms=300
        )
        self._validator_runs.append(run3)

        # Validate Date
        run4 = attach_line_validator(
            self.date_entry,
            lambda text: bool(text.strip()),
            "Date required",
            update_validation,
            is_visual=lambda: self._show_validation,
            debounce_ms=300
        )
        self._validator_runs.append(run4)

        # Validate Transmittal
        run5 = attach_line_validator(
            self.transmittal_entry,
            lambda text: bool(text.strip()),
            "Transmittal required",
            update_validation,
            is_visual=lambda: self._show_validation,
            debounce_ms=300
        )
        self._validator_runs.append(run5)

        # Validate Title
        run6 = attach_line_validator(
            self.from_title_entry,
            lambda text: bool(text.strip()),
            "Title required",
            update_validation,
            is_visual=lambda: self._show_validation,
            debounce_ms=300
        )
        self._validator_runs.append(run6)

        # Phone is optional - no validator needed

        # Validate Template Path
        run7 = attach_line_validator(
            self.template_entry,
            lambda text: bool(text.strip()) and os.path.exists(text.strip()),
            "Valid template file required",
            update_validation,
            is_visual=lambda: self._show_validation,
            debounce_ms=300
        )
        self._validator_runs.append(run7)

        # Validate Index Path
        run8 = attach_line_validator(
            self.index_entry,
            lambda text: bool(text.strip()) and os.path.exists(text.strip()),
            "Valid index file required",
            update_validation,
            is_visual=lambda: self._show_validation,
            debounce_ms=300
        )
        self._validator_runs.append(run8)

        # Validate Output Path
        run9 = attach_line_validator(
            self.output_entry,
            lambda text: bool(text.strip()) and os.path.exists(text.strip()),
            "Valid output folder required",
            update_validation,
            is_visual=lambda: self._show_validation,
            debounce_ms=300
        )
        self._validator_runs.append(run9)

        # Connect dropdown changes to validation
        self.pe_combo.currentTextChanged.connect(update_validation)
        self.firm_combo.currentTextChanged.connect(update_validation)

        # Initial validation (computes state but doesn't show red)
        update_validation()

    def _is_valid_email(self):
        """Check if from email is valid."""
        email = self.from_email_entry.text().strip()
        return is_valid_email(email) and bool(email)

    def _is_valid_project_name(self):
        """Check if project name is valid."""
        return bool(self.project_name_entry.text().strip())

    def _is_valid_template(self):
        """Check if template path is valid."""
        path = self.template_entry.text().strip()
        return bool(path) and os.path.exists(path)

    def _is_valid_index(self):
        """Check if index path is valid."""
        path = self.index_entry.text().strip()
        return bool(path) and os.path.exists(path)

    def _is_valid_output(self):
        """Check if output path is valid."""
        path = self.output_entry.text().strip()
        return bool(path) and os.path.exists(path)

    def _is_valid_cid_folder(self):
        """Check if CID folder path is valid."""
        path = self.cid_entry.text().strip()
        return bool(path) and os.path.exists(path)

    def _has_cid_documents(self):
        """Check if CID folder has any .cid files and document entries."""
        if not self._is_valid_cid_folder():
            return False

        # Check if there are any CID document rows with descriptions
        for row_data in self.cid_document_rows:
            if row_data['description'].text().strip():
                return True
        return False

    def _on_contact_field_changed(self):
        """Called when any contact field changes."""
        if self._show_validation:
            self._update_contact_borders()

    def _update_contact_borders(self):
        """Update red borders on contact fields."""
        from utils.validation import _mark_invalid, _mark_valid

        # Only update if validation is visible
        if not self._show_validation:
            return

        # Update each contact row
        for contact_data in self.contact_rows:
            name = contact_data['name'].text().strip()
            company = contact_data['company'].text().strip()
            email = contact_data['email'].text().strip()
            phone = contact_data['phone'].text().strip()

            # Mark name field
            if not name:
                _mark_invalid(contact_data['name'], "Name required")
            else:
                _mark_valid(contact_data['name'])

            # Mark company field (NOW REQUIRED)
            if not company:
                _mark_invalid(contact_data['company'], "Company required")
            else:
                _mark_valid(contact_data['company'])

            # Mark email field
            if not email:
                _mark_invalid(contact_data['email'], "Email required")
            else:
                _mark_valid(contact_data['email'])

            # Mark phone field (NOW REQUIRED)
            if not phone:
                _mark_invalid(contact_data['phone'], "Phone required")
            else:
                _mark_valid(contact_data['phone'])

    def _update_field_borders(self):
        """Update red borders on all fields based on current validation state."""
        from utils.validation import _mark_invalid, _mark_valid

        invalid = self._invalid_fields()
        invalid_widgets = {widget for _, widget in invalid if widget is not None}

        # Mark invalid fields (skip None widgets)
        for name, widget in invalid:
            if widget is not None:
                _mark_invalid(widget, f"{name} required")

        # Mark valid fields (remove red borders)
        all_fields = [
            self.template_entry, self.index_entry, self.output_entry,
            self.project_name_entry, self.project_number_entry, self.date_entry,
            self.transmittal_entry, self.from_title_entry, self.from_email_entry,
            self.pe_combo, self.firm_combo, self.description_text
            # Phone is optional - not validated
        ]
        for widget in all_fields:
            if widget not in invalid_widgets:
                _mark_valid(widget)

        # Update contact field borders
        self._update_contact_borders()

    def _reveal_validation(self):
        """Reveal validation errors (show red borders)."""
        if not self._show_validation:
            self._show_validation = True
            for run in self._validator_runs:
                run()  # Repaint with red where needed

        # Update all field borders
        self._update_field_borders()

    def _invalid_fields(self):
        """Get list of invalid fields with their widgets."""
        invalid = []

        # File Selection (mode-dependent)
        if not self._is_valid_template():
            invalid.append(("Template file", self.template_entry))

        if self.transmittal_type == "cid":
            # CID mode validation
            if not self._is_valid_cid_folder():
                invalid.append(("CID folder", self.cid_entry))
            if not self._has_cid_documents():
                invalid.append(("CID documents", self.cid_entry))
        else:
            # Standard mode validation
            if not self._is_valid_index():
                invalid.append(("Drawing index", self.index_entry))
            if not self._is_valid_output():
                invalid.append(("Output folder", self.output_entry))

        # Project Information
        if not self._is_valid_project_name():
            invalid.append(("Project name", self.project_name_entry))
        if not self.project_number_entry.text().strip():
            invalid.append(("Project number", self.project_number_entry))
        if not self.date_entry.text().strip():
            invalid.append(("Date", self.date_entry))
        if not self.transmittal_entry.text().strip():
            invalid.append(("Transmittal", self.transmittal_entry))

        # From Information
        if not self.from_name_entry.text().strip():
            invalid.append(("PE selection", self.pe_combo))
        if not self.from_title_entry.text().strip():
            invalid.append(("Title", self.from_title_entry))
        if not self._is_valid_email():
            invalid.append(("Email", self.from_email_entry))
        # Phone is optional - don't validate

        # Firm Number (check if a valid selection is made)
        firm_selection = self.firm_combo.currentText().strip()
        if not firm_selection:
            invalid.append(("Firm number", self.firm_combo))

        # Description
        if not self.description_text.toPlainText().strip():
            invalid.append(("Description", self.description_text))

        # Contacts - check if we have at least one valid contact
        # Don't mark the container, individual fields will be marked
        contacts = self.collect_contacts()
        if not contacts:
            # Add a message but don't mark a widget (individual fields will be marked)
            invalid.append(("At least one contact", None))

        return invalid

    def _focus_and_reveal(self, widget):
        """Reveal validation and focus on a specific widget."""
        self._reveal_validation()
        try:
            widget.setFocus()
            # Scroll to widget if in scroll area
            if hasattr(self, 'main_scroll'):
                self.main_scroll.ensureWidgetVisible(widget, 50, 50)
        except Exception:
            pass

    def auto_save(self):
        """Auto-save project data to JSON."""
        try:
            save_path = os.path.join(os.path.expanduser("~"), ".transmittal_builder_autosave.json")

            # Collect all data
            data = {
                "transmittal_type": self.transmittal_type,
                "template_path": self.template_entry.text(),
                "index_path": self.index_entry.text(),
                "output_folder": self.output_entry.text(),
                "cid_folder": self.cid_entry.text() if hasattr(self, 'cid_entry') else "",
                "cid_documents": self.collect_cid_index_data() if hasattr(self, 'cid_document_rows') else [],
                "fields": self.collect_fields(),
                "checks": self.collect_checks(),
                "contacts": self.collect_contacts(),
                "selected_pe": self.pe_combo.currentText(),  # Save PE selection
                "last_template_dir": self.last_template_dir,
                "last_index_dir": self.last_index_dir,
                "last_output_dir": self.last_output_dir,
                "timestamp": datetime.now().isoformat()
            }

            # Save to JSON
            with open(save_path, 'w') as f:
                json.dump(data, f, indent=2)

            # Don't log every auto-save to avoid spam
            # self.log_message("Project auto-saved", "info")

        except Exception:
            # Silently fail - don't spam the log
            pass

    def load_autosave(self):
        """Load auto-saved project data."""
        try:
            save_path = os.path.join(os.path.expanduser("~"), ".transmittal_builder_autosave.json")

            if not os.path.exists(save_path):
                return

            with open(save_path, 'r') as f:
                data = json.load(f)

            # Restore transmittal type
            transmittal_type = data.get("transmittal_type", "standard")
            if transmittal_type == "cid":
                self.cid_radio.setChecked(True)
                self.standard_radio.setChecked(False)
                self.transmittal_type = "cid"
            else:
                self.standard_radio.setChecked(True)
                self.cid_radio.setChecked(False)
                self.transmittal_type = "standard"

            # Update UI visibility
            self.update_file_section_visibility()

            # Restore file paths
            if data.get("template_path"):
                self.template_entry.setText(data["template_path"])
            if data.get("index_path"):
                self.index_entry.setText(data["index_path"])
            if data.get("output_folder"):
                self.output_entry.setText(data["output_folder"])
            if data.get("cid_folder") and hasattr(self, 'cid_entry'):
                self.cid_entry.setText(data["cid_folder"])

            # Restore directory memory
            self.last_template_dir = data.get("last_template_dir", "")
            self.last_index_dir = data.get("last_index_dir", "")
            self.last_output_dir = data.get("last_output_dir", "")

            # Restore fields
            fields = data.get("fields", {})
            if fields.get("date"):
                self.date_entry.setText(fields["date"])
            if fields.get("job_num"):
                self.project_number_entry.setText(fields["job_num"])
            if fields.get("transmittal_num") and hasattr(self, 'transmittal_entry'):
                self.transmittal_entry.setText(fields["transmittal_num"])
            if fields.get("client"):
                self.project_name_entry.setText(fields["client"])
            if fields.get("project_desc"):
                self.description_text.setPlainText(fields["project_desc"])

            # Restore PE selection (this will automatically populate from fields)
            selected_pe = data.get("selected_pe", "")
            if selected_pe:
                index = self.pe_combo.findText(selected_pe)
                if index >= 0:
                    self.pe_combo.setCurrentIndex(index)
                # The on_pe_changed handler will populate the from fields automatically

            if fields.get("firm"):
                index = self.firm_combo.findText(fields["firm"])
                if index >= 0:
                    self.firm_combo.setCurrentIndex(index)

            # Restore checkboxes
            checks = data.get("checks", {})
            for key, value in checks.items():
                checkbox = getattr(self, f"chk_{key}", None)
                if checkbox:
                    checkbox.setChecked(value)

            # Restore contacts (suppress logging during restoration)
            contacts = data.get("contacts", [])
            self._suppress_contact_log = True

            # Clear existing default contacts first
            while len(self.contact_rows) > 0:
                if self.contact_rows:
                    self.remove_contact_row(self.contact_rows[0]['widget'])

            # Add saved contacts
            for contact in contacts:
                self.add_contact_row()
                if self.contact_rows:
                    last_contact = self.contact_rows[-1]
                    last_contact['name'].setText(contact.get('name', ''))
                    last_contact['company'].setText(contact.get('company', ''))
                    last_contact['email'].setText(contact.get('email', ''))
                    last_contact['phone'].setText(contact.get('phone', ''))

            self._suppress_contact_log = False

            # Restore CID documents if in CID mode
            if self.transmittal_type == "cid" and hasattr(self, 'cid_document_rows'):
                cid_documents = data.get("cid_documents", [])
                if cid_documents:
                    # Clear existing CID rows
                    self.clear_cid_document_rows()

                    # Restore CID document entries
                    for cid_doc in cid_documents:
                        filename = cid_doc.get('filename', '')
                        description = cid_doc.get('description', '')
                        revision = cid_doc.get('revision', '-')

                        if filename:
                            self.add_cid_document_row(filename)
                            # Update the last added row
                            if self.cid_document_rows:
                                last_row = self.cid_document_rows[-1]
                                last_row['description'].setText(description)
                                last_row['revision'].setCurrentText(revision)

            # Log restoration without contact count
            self.log_message("Previous session restored", "success")

        except Exception as e:
            self.log_message(f"Failed to load auto-save: {str(e)}", "warning")

    def new_session(self):
        """Clear all fields and start a new session."""
        # Confirm with user
        reply = QMessageBox.question(
            self,
            "New Session",
            "Are you sure you want to start a new session?\n\nAll current data will be cleared.",
            QMessageBox.StandardButton.Yes | QMessageBox.StandardButton.No,
            QMessageBox.StandardButton.No
        )

        if reply == QMessageBox.StandardButton.Yes:
            # Clear file paths
            self.template_entry.clear()
            self.index_entry.clear()
            self.output_entry.clear()

            # Clear project info
            self.project_name_entry.clear()
            self.project_number_entry.clear()
            self.date_entry.setText(datetime.now().strftime("%m/%d/%Y"))  # Reset to today
            self.transmittal_entry.clear()
            self.description_text.clear()

            # Clear from info
            self.pe_combo.setCurrentIndex(0)
            self.from_title_entry.clear()
            self.from_email_entry.clear()
            self.from_phone_entry.clear()
            self.firm_combo.setCurrentIndex(0)

            # Clear all contacts (suppress logging)
            self._suppress_contact_log = True
            while len(self.contact_rows) > 0:
                if self.contact_rows:
                    self.remove_contact_row(self.contact_rows[0]['widget'])

            # Add one empty contact (suppress logging)
            self.add_contact_row()
            self._suppress_contact_log = False

            # Uncheck all checkboxes
            for checkbox_name in ['trans_pdf', 'trans_cad', 'trans_originals',
                                 'via_email', 'via_ftp',
                                 'ci_approval', 'ci_bid', 'ci_construction',
                                 'ci_asbuilt', 'ci_reference', 'ci_preliminary', 'ci_info',
                                 'vr_approved', 'vr_approved_noted', 'vr_rejected']:
                checkbox = getattr(self, f'chk_{checkbox_name}', None)
                if checkbox:
                    checkbox.setChecked(False)

            # Reset validation state
            self._show_validation = False
            from utils.validation import _mark_valid

            # Clear all red borders
            all_fields = [
                self.template_entry, self.index_entry, self.output_entry,
                self.project_name_entry, self.project_number_entry, self.date_entry,
                self.transmittal_entry, self.from_title_entry, self.from_email_entry,
                self.from_phone_entry, self.pe_combo, self.firm_combo, self.description_text
            ]
            for widget in all_fields:
                _mark_valid(widget)

            # Clear contact borders
            for contact_data in self.contact_rows:
                _mark_valid(contact_data['name'])
                _mark_valid(contact_data['company'])
                _mark_valid(contact_data['email'])
                _mark_valid(contact_data['phone'])

            # Delete autosave file
            try:
                save_path = os.path.join(os.path.expanduser("~"), ".transmittal_builder_autosave.json")
                if os.path.exists(save_path):
                    os.remove(save_path)
            except Exception:
                pass

            self.log_message("New session started", "success")

    # ========================================================================
    # EVENT HANDLERS
    # ========================================================================

    def browse_template(self):
        """Browse for template file."""
        start_dir = self.last_template_dir if self.last_template_dir else os.path.expanduser("~")

        file_path, _ = QFileDialog.getOpenFileName(
            self,
            "Select Transmittal Template",
            start_dir,
            "Word Documents (*.docx);;All Files (*.*)"
        )
        if file_path:
            self.template_entry.setText(file_path)
            self.last_template_dir = os.path.dirname(file_path)
            self.log_message(f"Template selected: {os.path.basename(file_path)}", "success")
            self.trigger_auto_save()

    def browse_index(self):
        """Browse for index file."""
        start_dir = self.last_index_dir if self.last_index_dir else os.path.expanduser("~")

        file_path, _ = QFileDialog.getOpenFileName(
            self,
            "Select Drawing Index",
            start_dir,
            "Excel Files (*.xlsx *.xls);;All Files (*.*)"
        )
        if file_path:
            self.index_entry.setText(file_path)
            self.last_index_dir = os.path.dirname(file_path)
            self.log_message(f"Index selected: {os.path.basename(file_path)}", "success")
            self.trigger_auto_save()

    def browse_output(self):
        """Browse for output folder."""
        start_dir = self.last_output_dir if self.last_output_dir else os.path.expanduser("~")

        folder_path = QFileDialog.getExistingDirectory(
            self,
            "Select Output Folder",
            start_dir
        )
        if folder_path:
            self.output_entry.setText(folder_path)
            self.last_output_dir = folder_path
            self.log_message(f"Output folder: {os.path.basename(folder_path)}", "success")
            self.trigger_auto_save()

    def on_pe_changed(self, pe_name: str):
        """Handle PE selection change."""
        if not pe_name or pe_name.strip() == "":
            # Clear fields if no PE selected
            self.from_name_entry.setText("")
            self.from_title_entry.setText("")
            self.from_email_entry.setText("")
            self.from_phone_entry.setText("")
            self.log_message("PE selection cleared", "info")
        else:
            # Get PE profile from config
            pe_info = config_manager.get_pe_profile(pe_name) if config_manager else None

            if pe_info:
                # Use profile data if available
                self.from_name_entry.setText(pe_info.get("name", pe_name))
                self.from_title_entry.setText(pe_info.get("title", ""))
                self.from_email_entry.setText(pe_info.get("email", ""))
                self.from_phone_entry.setText(pe_info.get("phone", ""))
                self.log_message(f"PE profile loaded: {pe_name}", "success")
            else:
                # No profile configured, just use the PE name
                self.from_name_entry.setText(pe_name)
                self.from_title_entry.setText("")  # Clear other fields
                self.from_email_entry.setText("")
                self.from_phone_entry.setText("")
                self.log_message(f"PE selected: {pe_name} (no profile configured)", "info")

        self.trigger_auto_save()

    def preview_transmittal(self):
        """Preview transmittal document."""
        self.log_message("Preview clicked", "info")

        # Check for invalid fields (pristine mode)
        bad = self._invalid_fields()
        if bad:
            self._focus_and_reveal(bad[0][1])
            bullets = "\n".join(f"• {name}" for name, _ in bad)
            show_warning_message(self, "Missing Information",
                               "Please complete the following before preview:\n\n" + bullets)
            return

        try:
            self.update_status("Generating preview...", OBSIDIAN_WARNING)

            # Create temporary file for preview
            import tempfile
            with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as tmp:
                preview_path = tmp.name

            # Handle different transmittal modes
            if self.transmittal_type == "cid":
                # CID mode: use CID folder and generate index from table
                cid_folder = self.cid_entry.text()
                cid_index_data = self.collect_cid_index_data()

                self.log_message("Rendering CID preview document...", "info")
                render_cid_transmittal(
                    self.template_entry.text(),
                    cid_folder,
                    cid_index_data,
                    self.collect_fields(),
                    self.collect_checks(),
                    self.collect_contacts(),
                    preview_path
                )
            else:
                # Standard mode: use PDF folder and Excel index
                output_folder = self.output_entry.text()

                self.log_message("Rendering preview document...", "info")
                render_transmittal(
                    self.template_entry.text(),
                    output_folder,  # documents_source (folder path)
                    self.index_entry.text(),
                    self.collect_fields(),
                    self.collect_checks(),
                    self.collect_contacts(),
                    preview_path,
                    None  # selected_files (None = use all files in folder)
                )

            self.log_message("Preview generated successfully", "success")
            self.update_status("Ready", OBSIDIAN_SUCCESS)

            # Open preview in default application
            self.log_message("Opening preview...", "info")
            if platform.system() == 'Windows':
                os.startfile(preview_path)
            elif platform.system() == 'Darwin':  # macOS
                subprocess.run(['open', preview_path])
            else:  # Linux
                subprocess.run(['xdg-open', preview_path])

            show_success_message(self, "Preview Ready",
                               "Preview document has been generated and opened.")

        except Exception as e:
            self.log_message(f"Preview failed: {str(e)}", "error")
            self.update_status("Error", OBSIDIAN_ERROR)
            show_error_message(self, "Preview Error",
                             f"Failed to generate preview:\n\n{str(e)}")

    def generate_transmittal(self):
        """Generate transmittal document."""
        self.log_message("Generate clicked", "info")

        # Check for invalid fields (pristine mode)
        bad = self._invalid_fields()
        if bad:
            self._focus_and_reveal(bad[0][1])
            bullets = "\n".join(f"• {name}" for name, _ in bad)
            show_warning_message(self, "Missing Information",
                               "Please complete the following before generating:\n\n" + bullets)
            return

        try:
            self.update_status("Generating transmittal...", OBSIDIAN_WARNING)

            # Create output filename
            project_num = self.project_number_entry.text().strip()
            timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")

            # Handle different transmittal modes
            if self.transmittal_type == "cid":
                # CID mode: save to CID folder
                cid_folder = self.cid_entry.text()
                filename = f"CID_Transmittal_{project_num}_{timestamp}.docx"
                output_path = os.path.join(cid_folder, filename)
                cid_index_data = self.collect_cid_index_data()

                self.log_message("Rendering CID transmittal document...", "info")
                render_cid_transmittal(
                    self.template_entry.text(),
                    cid_folder,
                    cid_index_data,
                    self.collect_fields(),
                    self.collect_checks(),
                    self.collect_contacts(),
                    output_path
                )
            else:
                # Standard mode: save to PDF folder
                output_folder = self.output_entry.text()
                filename = f"Transmittal_{project_num}_{timestamp}.docx"
                output_path = os.path.join(output_folder, filename)

                self.log_message("Rendering transmittal document...", "info")
                render_transmittal(
                    self.template_entry.text(),
                    output_folder,  # documents_source (folder path)
                    self.index_entry.text(),
                    self.collect_fields(),
                    self.collect_checks(),
                    self.collect_contacts(),
                    output_path,
                    None  # selected_files (None = use all files in folder)
                )

            self.log_message(f"Transmittal saved: {filename}", "success")
            self.update_status("Ready", OBSIDIAN_SUCCESS)

            # Ask if user wants to open the file
            if show_question_dialog(self, "Success",
                                   f"Transmittal generated successfully!\n\n{filename}\n\nWould you like to open it?"):
                self.log_message("Opening transmittal...", "info")
                if platform.system() == 'Windows':
                    os.startfile(output_path)
                elif platform.system() == 'Darwin':  # macOS
                    subprocess.run(['open', output_path])
                else:  # Linux
                    subprocess.run(['xdg-open', output_path])

            show_success_message(self, "Generation Complete",
                               f"Transmittal has been generated:\n\n{filename}")

        except Exception as e:
            self.log_message(f"Generation failed: {str(e)}", "error")
            self.update_status("Error", OBSIDIAN_ERROR)
            show_error_message(self, "Generation Error",
                             f"Failed to generate transmittal:\n\n{str(e)}")

    def open_bug_report(self):
        """Open bug report dialog."""
        self.log_message("Bug report opened", "info")

        # Get session log from activity log
        session_log = self.log_textbox.toPlainText()

        # Open dialog
        dialog = BugReportDialog(self, session_log)
        if dialog.exec() == QDialog.DialogCode.Accepted:
            self.log_message("Bug report sent", "success")
            show_success_message(self, "Thank You!",
                               "Your bug report has been sent. We'll look into it!")

    def open_suggestion(self):
        """Open suggestion dialog."""
        self.log_message("Suggestion opened", "info")

        # Open dialog
        dialog = SuggestionDialog(self)
        if dialog.exec() == QDialog.DialogCode.Accepted:
            self.log_message("Suggestion sent", "success")
            show_success_message(self, "Thank You!",
                               "Your suggestion has been sent. We appreciate your feedback!")

    def open_settings(self):
        """Open settings (built-in YAML editor)."""
        self.log_message("Opening settings...", "info")

        try:
            if config_manager:
                config_path = config_manager.get_config_path()
                self.log_message(f"Config file: {config_path}", "info")

                # Open built-in YAML editor
                dialog = YAMLEditorDialog(config_path, self)
                if dialog.exec() == QDialog.DialogCode.Accepted:
                    self.log_message("Settings saved", "success")
                    # Optionally reload config here (requires restart for now)
                else:
                    self.log_message("Settings editor closed", "info")
            else:
                show_error_message(self, "Settings Error",
                                 "Config system not available")
        except Exception as e:
            self.log_message(f"Failed to open settings: {e}", "error")
            show_error_message(self, "Settings Error",
                             f"Failed to open settings:\n\n{str(e)}")

# ============================================================================
# DIALOG CLASSES
# ============================================================================

class BugReportDialog(QDialog):
    """Dialog for submitting bug reports."""

    def __init__(self, parent=None, session_log: str = ""):
        super().__init__(parent)
        self.session_log = session_log
        self.setWindowTitle("Report a Bug")
        self.setModal(True)
        self.setMinimumSize(700, 600)

        # Apply obsidian theme
        self.setStyleSheet(f"""
            QDialog {{
                background-color: {OBSIDIAN_BG};
            }}
        """)

        self.init_ui()

    def init_ui(self):
        """Initialize the dialog UI."""
        layout = QVBoxLayout(self)
        layout.setContentsMargins(20, 20, 20, 20)
        layout.setSpacing(15)

        # Title
        title_label = QLabel("🐛 Report a Bug")
        title_label.setStyleSheet(f"""
            font-size: 18px;
            font-weight: bold;
            color: {OBSIDIAN_TEXT_PRIMARY};
            padding: 10px;
        """)
        layout.addWidget(title_label)

        # Description
        desc_label = QLabel("Please describe the bug you encountered:")
        desc_label.setStyleSheet(f"color: {OBSIDIAN_TEXT_SECONDARY};")
        layout.addWidget(desc_label)

        # Bug description text area
        self.bug_description = QTextEdit()
        self.bug_description.setPlaceholderText("Describe what happened, what you expected, and steps to reproduce...")
        self.bug_description.setMinimumHeight(150)
        layout.addWidget(self.bug_description)

        # Email field
        email_layout = QHBoxLayout()
        email_label = QLabel("Your Email (optional):")
        email_label.setFixedWidth(150)
        email_label.setStyleSheet(f"color: {OBSIDIAN_TEXT_SECONDARY};")
        self.email_entry = QLineEdit()
        self.email_entry.setPlaceholderText("your.email@example.com")
        email_layout.addWidget(email_label)
        email_layout.addWidget(self.email_entry)
        layout.addLayout(email_layout)

        # Session log section
        log_label = QLabel("Session Log (automatically included):")
        log_label.setStyleSheet(f"color: {OBSIDIAN_TEXT_SECONDARY}; margin-top: 10px;")
        layout.addWidget(log_label)

        # Session log display
        self.log_display = QPlainTextEdit()
        self.log_display.setPlainText(self.session_log)
        self.log_display.setReadOnly(True)
        self.log_display.setMaximumHeight(150)
        self.log_display.setStyleSheet(f"""
            QPlainTextEdit {{
                background-color: {OBSIDIAN_CARD_DEEP};
                color: {OBSIDIAN_TEXT_SECONDARY};
                border: 1px solid {OBSIDIAN_BORDER};
                border-radius: 6px;
                padding: 8px;
                font-family: 'Consolas', 'Courier New', monospace;
                font-size: 10px;
            }}
        """)
        layout.addWidget(self.log_display)

        # Buttons
        button_layout = QHBoxLayout()
        button_layout.addStretch()

        cancel_btn = QPushButton("Cancel")
        cancel_btn.setFixedWidth(100)
        cancel_btn.setCursor(Qt.CursorShape.PointingHandCursor)
        cancel_btn.clicked.connect(self.reject)
        button_layout.addWidget(cancel_btn)

        send_btn = QPushButton("Send Report")
        send_btn.setFixedWidth(120)
        send_btn.setCursor(Qt.CursorShape.PointingHandCursor)
        send_btn.setStyleSheet(f"""
            QPushButton {{
                background: qlineargradient(
                    x1:0, y1:0, x2:0, y2:1,
                    stop:0 {ROOT3_BLUE},
                    stop:1 #0F4A7A
                );
                color: {OBSIDIAN_TEXT_PRIMARY};
                border: 2px solid {ROOT3_BLUE};
                border-radius: 6px;
                padding: 8px;
                font-weight: bold;
            }}
            QPushButton:hover {{
                background: qlineargradient(
                    x1:0, y1:0, x2:0, y2:1,
                    stop:0 #2E7BC6,
                    stop:1 {ROOT3_BLUE}
                );
            }}
        """)
        send_btn.clicked.connect(self.send_report)
        button_layout.addWidget(send_btn)

        layout.addLayout(button_layout)

    def keyPressEvent(self, event):
        """Override to prevent Enter from closing dialog."""
        # Only close on Ctrl+Enter, not just Enter
        if event.key() == Qt.Key.Key_Return or event.key() == Qt.Key.Key_Enter:
            if event.modifiers() == Qt.KeyboardModifier.ControlModifier:
                self.send_report()
            # Otherwise, ignore the event (don't close dialog)
            return
        super().keyPressEvent(event)

    def get_system_info(self) -> str:
        """Get system information."""
        import platform
        return f"{platform.system()} {platform.release()} | Python {platform.python_version()} | PyQt6"

    def send_report(self):
        """Send the bug report via email using new email system."""
        description = self.bug_description.toPlainText().strip()

        if not description:
            QMessageBox.warning(self, "Missing Information",
                              "Please describe the bug before sending.")
            return

        user_email = self.email_entry.text().strip()

        # Show sending message
        from PyQt6.QtWidgets import QProgressDialog
        progress = QProgressDialog("Sending bug report...", None, 0, 0, self)
        progress.setWindowTitle("Sending")
        progress.setWindowModality(Qt.WindowModality.WindowModal)
        progress.setCancelButton(None)
        progress.setMinimumDuration(0)
        progress.setValue(0)
        progress.show()
        QApplication.processEvents()

        try:
            # Prepare payload for new email system (simplified)
            payload = {
                "user_email": user_email or "N/A",
                "version": VERSION,
                "submitted": None,  # Will use current time
                "description": description,
                "session_log": self.session_log
            }

            # Get email credentials from config
            mail_creds = config_manager.mail_creds()

            # Send using new email system
            logo_path = LOGO_PATH if os.path.exists(LOGO_PATH) else ""
            success, message = send_bug_report_email(
                payload,
                to=mail_creds.default_receiver,
                cc=user_email if user_email else None,
                logo_path=logo_path,
                sender=mail_creds.sender,
                password=mail_creds.app_password
            )

            progress.close()

            if success:
                # Just close the dialog - no second popup
                self.accept()
            else:
                show_error_message(self, "Send Failed",
                                 f"Failed to send bug report:\n\n{message}")
        except Exception as e:
            progress.close()
            show_error_message(self, "Error", f"Unexpected error: {str(e)}")


class SuggestionDialog(QDialog):
    """Dialog for submitting feature suggestions."""

    def __init__(self, parent=None):
        super().__init__(parent)
        self.setWindowTitle("Submit a Suggestion")
        self.setModal(True)
        self.setMinimumSize(600, 500)

        # Apply obsidian theme
        self.setStyleSheet(f"""
            QDialog {{
                background-color: {OBSIDIAN_BG};
            }}
        """)

        self.init_ui()

    def init_ui(self):
        """Initialize the dialog UI."""
        layout = QVBoxLayout(self)
        layout.setContentsMargins(20, 20, 20, 20)
        layout.setSpacing(15)

        # Title
        title_label = QLabel("💡 Submit a Suggestion")
        title_label.setStyleSheet(f"""
            font-size: 18px;
            font-weight: bold;
            color: {OBSIDIAN_TEXT_PRIMARY};
            padding: 10px;
        """)
        layout.addWidget(title_label)

        # Description
        desc_label = QLabel("We'd love to hear your ideas for improving the Transmittal Builder!")
        desc_label.setStyleSheet(f"color: {OBSIDIAN_TEXT_SECONDARY};")
        desc_label.setWordWrap(True)
        layout.addWidget(desc_label)

        # Category
        category_layout = QHBoxLayout()
        category_label = QLabel("Category:")
        category_label.setFixedWidth(100)
        category_label.setStyleSheet(f"color: {OBSIDIAN_TEXT_SECONDARY};")
        self.category_combo = QComboBox()
        self.category_combo.addItems([
            "Feature Request",
            "UI/UX Improvement",
            "Performance",
            "Documentation",
            "Other"
        ])
        category_layout.addWidget(category_label)
        category_layout.addWidget(self.category_combo)
        layout.addLayout(category_layout)

        # Suggestion text area
        suggestion_label = QLabel("Your Suggestion:")
        suggestion_label.setStyleSheet(f"color: {OBSIDIAN_TEXT_SECONDARY};")
        layout.addWidget(suggestion_label)

        self.suggestion_text = QTextEdit()
        self.suggestion_text.setPlaceholderText("Describe your suggestion in detail...")
        self.suggestion_text.setMinimumHeight(200)
        layout.addWidget(self.suggestion_text)

        # Email field
        email_layout = QHBoxLayout()
        email_label = QLabel("Your Email (optional):")
        email_label.setFixedWidth(150)
        email_label.setStyleSheet(f"color: {OBSIDIAN_TEXT_SECONDARY};")
        self.email_entry = QLineEdit()
        self.email_entry.setPlaceholderText("your.email@example.com")
        email_layout.addWidget(email_label)
        email_layout.addWidget(self.email_entry)
        layout.addLayout(email_layout)

        # Buttons
        button_layout = QHBoxLayout()
        button_layout.addStretch()

        cancel_btn = QPushButton("Cancel")
        cancel_btn.setFixedWidth(100)
        cancel_btn.setCursor(Qt.CursorShape.PointingHandCursor)
        cancel_btn.clicked.connect(self.reject)
        button_layout.addWidget(cancel_btn)

        send_btn = QPushButton("Send Suggestion")
        send_btn.setFixedWidth(140)
        send_btn.setCursor(Qt.CursorShape.PointingHandCursor)
        send_btn.setStyleSheet(f"""
            QPushButton {{
                background: qlineargradient(
                    x1:0, y1:0, x2:0, y2:1,
                    stop:0 {ROOT3_BLUE},
                    stop:1 #0F4A7A
                );
                color: {OBSIDIAN_TEXT_PRIMARY};
                border: 2px solid {ROOT3_BLUE};
                border-radius: 6px;
                padding: 8px;
                font-weight: bold;
            }}
            QPushButton:hover {{
                background: qlineargradient(
                    x1:0, y1:0, x2:0, y2:1,
                    stop:0 #2E7BC6,
                    stop:1 {ROOT3_BLUE}
                );
            }}
        """)
        send_btn.clicked.connect(self.send_suggestion)
        button_layout.addWidget(send_btn)

        layout.addLayout(button_layout)

    def keyPressEvent(self, event):
        """Override to prevent Enter from closing dialog."""
        # Only close on Ctrl+Enter, not just Enter
        if event.key() == Qt.Key.Key_Return or event.key() == Qt.Key.Key_Enter:
            if event.modifiers() == Qt.KeyboardModifier.ControlModifier:
                self.send_suggestion()
            # Otherwise, ignore the event (don't close dialog)
            return
        super().keyPressEvent(event)

    def send_suggestion(self):
        """Send the suggestion via email using new email system."""
        suggestion = self.suggestion_text.toPlainText().strip()

        if not suggestion:
            QMessageBox.warning(self, "Missing Information",
                              "Please describe your suggestion before sending.")
            return

        category = self.category_combo.currentText()
        user_email = self.email_entry.text().strip()

        # Show sending message
        from PyQt6.QtWidgets import QProgressDialog
        progress = QProgressDialog("Sending suggestion...", None, 0, 0, self)
        progress.setWindowTitle("Sending")
        progress.setWindowModality(Qt.WindowModality.WindowModal)
        progress.setCancelButton(None)
        progress.setMinimumDuration(0)
        progress.setValue(0)
        progress.show()
        QApplication.processEvents()

        try:
            # Prepare payload for new email system
            payload = {
                "category": category,
                "user_email": user_email or "N/A",
                "submitted": None,  # Will use current time
                "suggestion": suggestion,
                "version": VERSION
            }

            # Get email credentials from config
            mail_creds = config_manager.mail_creds()

            # Send using new email system
            logo_path = LOGO_PATH if os.path.exists(LOGO_PATH) else ""
            success, message = send_suggestion_email(
                payload,
                to=mail_creds.default_receiver,
                cc=user_email if user_email else None,
                logo_path=logo_path,
                sender=mail_creds.sender,
                password=mail_creds.app_password
            )

            progress.close()

            if success:
                # Just close the dialog - no second popup
                self.accept()
            else:
                show_error_message(self, "Send Failed",
                                 f"Failed to send suggestion:\n\n{message}")
        except Exception as e:
            progress.close()
            show_error_message(self, "Error", f"Unexpected error: {str(e)}")


# ============================================================================
# CUSTOM MESSAGE BOX FUNCTIONS
# ============================================================================

def show_info_message(parent, title: str, message: str):
    """Show an info message with obsidian theme."""
    msg_box = QMessageBox(parent)
    msg_box.setWindowTitle(title)
    msg_box.setText(message)
    msg_box.setIcon(QMessageBox.Icon.Information)
    msg_box.setStyleSheet(f"""
        QMessageBox {{
            background-color: {OBSIDIAN_BG};
        }}
        QMessageBox QLabel {{
            color: {OBSIDIAN_TEXT_PRIMARY};
        }}
        QPushButton {{
            background: qlineargradient(
                x1:0, y1:0, x2:0, y2:1,
                stop:0 {OBSIDIAN_BLUE_DEEP},
                stop:1 {OBSIDIAN_CARD_MID}
            );
            color: {OBSIDIAN_TEXT_PRIMARY};
            border: 2px solid {OBSIDIAN_BORDER};
            border-radius: 6px;
            padding: 6px 20px;
            min-width: 80px;
        }}
        QPushButton:hover {{
            border: 2px solid {OBSIDIAN_BLUE_GLOW};
        }}
    """)
    msg_box.exec()


def show_warning_message(parent, title: str, message: str):
    """Show a warning message with obsidian theme."""
    msg_box = QMessageBox(parent)
    msg_box.setWindowTitle(title)
    msg_box.setText(message)
    msg_box.setIcon(QMessageBox.Icon.Warning)
    msg_box.setStyleSheet(f"""
        QMessageBox {{
            background-color: {OBSIDIAN_BG};
        }}
        QMessageBox QLabel {{
            color: {OBSIDIAN_TEXT_PRIMARY};
        }}
        QPushButton {{
            background: qlineargradient(
                x1:0, y1:0, x2:0, y2:1,
                stop:0 {OBSIDIAN_WARNING},
                stop:1 #CC8A3D
            );
            color: {OBSIDIAN_TEXT_PRIMARY};
            border: 2px solid {OBSIDIAN_WARNING};
            border-radius: 6px;
            padding: 6px 20px;
            min-width: 80px;
        }}
        QPushButton:hover {{
            background: qlineargradient(
                x1:0, y1:0, x2:0, y2:1,
                stop:0 #FFC85D,
                stop:1 {OBSIDIAN_WARNING}
            );
        }}
    """)
    msg_box.exec()


def show_error_message(parent, title: str, message: str):
    """Show an error message with obsidian theme."""
    msg_box = QMessageBox(parent)
    msg_box.setWindowTitle(title)
    msg_box.setText(message)
    msg_box.setIcon(QMessageBox.Icon.Critical)
    msg_box.setStyleSheet(f"""
        QMessageBox {{
            background-color: {OBSIDIAN_BG};
        }}
        QMessageBox QLabel {{
            color: {OBSIDIAN_TEXT_PRIMARY};
        }}
        QPushButton {{
            background: qlineargradient(
                x1:0, y1:0, x2:0, y2:1,
                stop:0 {OBSIDIAN_ERROR},
                stop:1 #CC4A5E
            );
            color: {OBSIDIAN_TEXT_PRIMARY};
            border: 2px solid {OBSIDIAN_ERROR};
            border-radius: 6px;
            padding: 6px 20px;
            min-width: 80px;
        }}
        QPushButton:hover {{
            background: qlineargradient(
                x1:0, y1:0, x2:0, y2:1,
                stop:0 #FF6A7E,
                stop:1 {OBSIDIAN_ERROR}
            );
        }}
    """)
    msg_box.exec()


def show_success_message(parent, title: str, message: str):
    """Show a success message with obsidian theme."""
    msg_box = QMessageBox(parent)
    msg_box.setWindowTitle(title)
    msg_box.setText(message)
    msg_box.setIcon(QMessageBox.Icon.Information)
    msg_box.setStyleSheet(f"""
        QMessageBox {{
            background-color: {OBSIDIAN_BG};
        }}
        QMessageBox QLabel {{
            color: {OBSIDIAN_TEXT_PRIMARY};
        }}
        QPushButton {{
            background: qlineargradient(
                x1:0, y1:0, x2:0, y2:1,
                stop:0 {OBSIDIAN_SUCCESS},
                stop:1 #3AAE70
            );
            color: {OBSIDIAN_TEXT_PRIMARY};
            border: 2px solid {OBSIDIAN_SUCCESS};
            border-radius: 6px;
            padding: 6px 20px;
            min-width: 80px;
        }}
        QPushButton:hover {{
            background: qlineargradient(
                x1:0, y1:0, x2:0, y2:1,
                stop:0 #5AEE90,
                stop:1 {OBSIDIAN_SUCCESS}
            );
        }}
    """)
    msg_box.exec()


def show_question_dialog(parent, title: str, message: str) -> bool:
    """Show a yes/no question dialog with obsidian theme. Returns True if Yes."""
    msg_box = QMessageBox(parent)
    msg_box.setWindowTitle(title)
    msg_box.setText(message)
    msg_box.setIcon(QMessageBox.Icon.Question)
    msg_box.setStandardButtons(QMessageBox.StandardButton.Yes | QMessageBox.StandardButton.No)
    msg_box.setDefaultButton(QMessageBox.StandardButton.No)
    msg_box.setStyleSheet(f"""
        QMessageBox {{
            background-color: {OBSIDIAN_BG};
        }}
        QMessageBox QLabel {{
            color: {OBSIDIAN_TEXT_PRIMARY};
        }}
        QPushButton {{
            background: qlineargradient(
                x1:0, y1:0, x2:0, y2:1,
                stop:0 {OBSIDIAN_BLUE_DEEP},
                stop:1 {OBSIDIAN_CARD_MID}
            );
            color: {OBSIDIAN_TEXT_PRIMARY};
            border: 2px solid {OBSIDIAN_BORDER};
            border-radius: 6px;
            padding: 6px 20px;
            min-width: 80px;
        }}
        QPushButton:hover {{
            border: 2px solid {OBSIDIAN_BLUE_GLOW};
        }}
    """)
    result = msg_box.exec()
    return result == QMessageBox.StandardButton.Yes


# ============================================================================
# MAIN ENTRY POINT
# ============================================================================

def main():
    """Main entry point for the application."""
    # In PyQt6, high DPI scaling is enabled by default
    # No need to set AA_EnableHighDpiScaling and AA_UseHighDpiPixmaps

    app = QApplication(sys.argv)

    # Set application-wide font with responsive size
    screen = app.primaryScreen()
    screen_width = screen.availableGeometry().width()

    # Responsive font sizing
    if screen_width <= 1366:
        font_size = 8  # Smaller font for small screens
    elif screen_width <= 1920:
        font_size = 9  # Standard font for medium screens
    else:
        font_size = 10  # Larger font for large screens

    font = QFont("Segoe UI", font_size)
    font.setStyleHint(QFont.StyleHint.SansSerif)
    app.setFont(font)

    # Create and show main window
    window = TransmittalBuilderPyQt6()
    window.show()

    sys.exit(app.exec())

if __name__ == "__main__":
    main()

